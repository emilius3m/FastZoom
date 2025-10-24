#!/usr/bin/env python3
"""
Script to create a proper Word template for Giornale di Cantiere export
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def create_giornale_template():
    """Create a Word template with all required placeholders for Giornale di Cantiere"""
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph('REGISTRO GIORNALE DI CANTIERE')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Add some spacing
    doc.add_paragraph('')
    
    # Section 1: Site Information
    doc.add_paragraph('INFORMAZIONI SITO').runs[0].bold = True
    doc.add_paragraph('Nome Sito: {{sito_nome}}')
    doc.add_paragraph('Codice Sito: {{sito_codice}}')
    doc.add_paragraph('Località: {{sito_localita}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 2: Export Metadata
    doc.add_paragraph('METADATI EXPORT').runs[0].bold = True
    doc.add_paragraph('Data Export: {{data_export}}')
    doc.add_paragraph('Utente Export: {{utente_export}}')
    doc.add_paragraph('Filtri Applicati: {{filtri_applicati}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 3: Statistics
    doc.add_paragraph('STATISTICHE RIEPILOGATIVE').runs[0].bold = True
    doc.add_paragraph('Totale Giornali: {{totale_giornali}}')
    doc.add_paragraph('Giornali Validati: {{giornali_validati}}')
    doc.add_paragraph('Giornali Pendenti: {{giornali_pendenti}}')
    doc.add_paragraph('Operatori Attivi: {{operatori_attivi}}')
    doc.add_paragraph('Percentuale Completamento: {{percentuale_completamento}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 4: Giornali List Table
    doc.add_paragraph('ELENCO GIORNALI').runs[0].bold = True
    
    # Create table for giornali list
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Data', 'Orari', 'Responsabile', 'Condizioni Meteo', 'Stato', 'Note']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Empty data row (will be populated dynamically)
    for i in range(6):
        table.cell(1, i).text = ''
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 5: Diary Information
    doc.add_paragraph('DETTAGLIO GIORNALE').runs[0].bold = True
    doc.add_paragraph('DATA GIORNALE: {{giornale_data}}')
    doc.add_paragraph('ORA INIZIO: {{giornale_ora_inizio}}')
    doc.add_paragraph('ORA FINE: {{giornale_ora_fine}}')
    doc.add_paragraph('COMPILATORE: {{giornale_compilatore}}')
    doc.add_paragraph('RESPONSABILE SCAVO: {{giornale_responsabile}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 6: Working Conditions
    doc.add_paragraph('CONDIZIONI OPERATIVE').runs[0].bold = True
    doc.add_paragraph('Condizioni Meteo: {{condizioni_meteo}}')
    doc.add_paragraph('Temperatura Minima: {{temperatura_min}}°C')
    doc.add_paragraph('Temperatura Massima: {{temperatura_max}}°C')
    doc.add_paragraph('Note Meteo: {{note_meteo}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 7: Work Description
    doc.add_paragraph('DESCRIZIONE LAVORI').runs[0].bold = True
    doc.add_paragraph('Descrizione Lavori: {{descrizione_lavori}}')
    doc.add_paragraph('')
    doc.add_paragraph('Modalità Lavorazioni: {{modalita_lavorazioni}}')
    doc.add_paragraph('')
    doc.add_paragraph('Attrezzatura Utilizzata: {{attrezzatura_utilizzata}}')
    doc.add_paragraph('')
    doc.add_paragraph('Mezzi Utilizzati: {{mezzi_utilizzati}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 8: Archaeological Documentation
    doc.add_paragraph('DOCUMENTAZIONE ARCHEOLOGICA').runs[0].bold = True
    doc.add_paragraph('US Elaborate: {{us_elaborate}}')
    doc.add_paragraph('USM Elaborate: {{usm_elaborate}}')
    doc.add_paragraph('USR Elaborate: {{usr_elaborate}}')
    doc.add_paragraph('Materiali Rinvenuti: {{materiali_rinvenuti}}')
    doc.add_paragraph('Documentazione Prodotta: {{documentazione_prodotta}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 9: Operators Table
    doc.add_paragraph('OPERATORI PRESENTI').runs[0].bold = True
    
    # Create table for operators
    operators_table = doc.add_table(rows=2, cols=3)
    operators_table.style = 'Table Grid'
    
    # Header row
    op_headers = ['Nome', 'Qualifica', 'Ruolo']
    for i, header in enumerate(op_headers):
        cell = operators_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Empty data row (will be populated dynamically)
    for i in range(3):
        operators_table.cell(1, i).text = ''
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 10: Inspections and Dispositions
    doc.add_paragraph('SOPRALLUOGHI E DISPOSIZIONI').runs[0].bold = True
    doc.add_paragraph('Sopralluoghi: {{sopralluoghi}}')
    doc.add_paragraph('Disposizioni RUP: {{disposizioni_rup}}')
    doc.add_paragraph('Disposizioni Direttore: {{disposizioni_direttore}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 11: Events and Issues
    doc.add_paragraph('EVENTI PARTICOLARI').runs[0].bold = True
    doc.add_paragraph('Contestazioni: {{contestazioni}}')
    doc.add_paragraph('Sospensioni: {{sospensioni}}')
    doc.add_paragraph('Incidenti: {{incidenti}}')
    doc.add_paragraph('Forniture: {{forniture}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 12: Notes and Issues
    doc.add_paragraph('NOTE E PROBLEMATICHE').runs[0].bold = True
    doc.add_paragraph('Note Generali: {{note_generali}}')
    doc.add_paragraph('Problematiche: {{problematiche}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 13: Validation
    doc.add_paragraph('STATO VALIDAZIONE').runs[0].bold = True
    doc.add_paragraph('Stato Validazione: {{stato_validazione}}')
    doc.add_paragraph('Data Validazione: {{data_validazione}}')
    
    # Add spacing
    doc.add_paragraph('')
    
    # Section 14: Signatures
    doc.add_paragraph('FIRME').runs[0].bold = True
    doc.add_paragraph('')
    
    # Create signature table
    sig_table = doc.add_table(rows=3, cols=1)
    sig_table.style = 'Table Grid'
    
    sig_table.cell(0, 0).text = '_________________________\nResponsabile Scavo'
    sig_table.cell(1, 0).text = '_________________________\nDirettore Lavori'
    sig_table.cell(2, 0).text = '_________________________\nRUP (Responsabile Unico Procedimento)'
    
    # Save the document
    template_path = 'app/templates/word/Giornale_Template_con_Placeholder.docx'
    doc.save(template_path)
    print(f"Template created successfully at: {template_path}")
    
    return template_path

if __name__ == "__main__":
    create_giornale_template()