# app/services/us_word_export.py
"""
Export US/USM compilando direttamente template Word esistente
Mantiene layout IDENTICO al documento MiC originale
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from pathlib import Path
from typing import Optional, Dict, Any
import re


class USWordExporter:
    """
    Compila template Word esistente con dati US
    NON crea nuovo template - modifica l'originale mantenendo il layout
    """

    def __init__(self, template_path: str):
        """
        Args:
            template_path: Path al file .docx template originale
        """
        self.template_path = template_path

    def export_us(self, us_data: Dict[str, Any], output_path: str) -> str:
        """
        Compila template US con dati database

        Args:
            us_data: Dizionario con dati US dal database
            output_path: Path dove salvare il documento compilato

        Returns:
            Path del file generato
        """
        # Carica template originale
        doc = Document(self.template_path)

        # COMPILA TUTTI I CAMPI DEL TEMPLATE
        # Cerca e sostituisci placeholder in tutto il documento

        # 1. INTESTAZIONE
        self._replace_text(doc, '{{us_code}}', us_data.get('us_code', ''))
        self._replace_text(doc, '{{ente_responsabile}}', us_data.get('ente_responsabile', ''))
        self._replace_text(doc, '{{anno}}', str(us_data.get('anno', '')))
        self._replace_text(doc, '{{identificativo_rif}}', us_data.get('identificativo_rif', ''))

        # 2. LOCALIZZAZIONE
        self._replace_text(doc, '{{localita}}', us_data.get('localita', ''))
        self._replace_text(doc, '{{area_struttura}}', us_data.get('area_struttura', ''))
        self._replace_text(doc, '{{ambiente}}', us_data.get('ambiente_unita_funzione', ''))
        self._replace_text(doc, '{{settori}}', us_data.get('settori', ''))

        # 3. DOCUMENTAZIONE - Riferimenti testuali (come nel template)
        self._replace_text(doc, '{{piante}}', us_data.get('piante_riferimenti', ''))
        self._replace_text(doc, '{{prospetti}}', us_data.get('prospetti_riferimenti', ''))
        self._replace_text(doc, '{{sezioni}}', us_data.get('sezioni_riferimenti', ''))

        # 4. DEFINIZIONE E CARATTERIZZAZIONE
        self._replace_text(doc, '{{definizione}}', us_data.get('definizione', ''))
        self._replace_text(doc, '{{criteri}}', us_data.get('criteri_distinzione', ''))
        self._replace_text(doc, '{{formazione}}', us_data.get('modo_formazione', ''))

        # 5. COMPONENTI
        self._replace_text(doc, '{{inorganici}}', us_data.get('componenti_inorganici', ''))
        self._replace_text(doc, '{{organici}}', us_data.get('componenti_organici', ''))
        self._replace_text(doc, '{{consistenza}}', us_data.get('consistenza', ''))
        self._replace_text(doc, '{{colore}}', us_data.get('colore', ''))
        self._replace_text(doc, '{{misure}}', us_data.get('misure', ''))
        self._replace_text(doc, '{{conservazione}}', us_data.get('stato_conservazione', ''))

        # 6. SEQUENZA FISICA - compila tabella relazioni
        self._compile_matrix_harris_table(doc, us_data.get('sequenza_fisica', {}))

        # 7. DESCRIZIONE
        self._replace_text(doc, '{{descrizione}}', us_data.get('descrizione', ''))
        self._replace_text(doc, '{{osservazioni}}', us_data.get('osservazioni', ''))
        self._replace_text(doc, '{{interpretazione}}', us_data.get('interpretazione', ''))

        # 8. DATAZIONE
        self._replace_text(doc, '{{datazione}}', us_data.get('datazione', ''))
        self._replace_text(doc, '{{periodo}}', us_data.get('periodo', ''))
        self._replace_text(doc, '{{fase}}', us_data.get('fase', ''))
        self._replace_text(doc, '{{elementi_datanti}}', us_data.get('elementi_datanti', ''))

        # 9. REPERTI
        self._replace_text(doc, '{{reperti}}', us_data.get('dati_quantitativi_reperti', ''))

        # 10. CAMPIONATURE - compila checkbox
        campionature = us_data.get('campionature', {})
        self._replace_checkbox(doc, '{{flottazione}}', campionature.get('flottazione', False))
        self._replace_checkbox(doc, '{{setacciatura}}', campionature.get('setacciatura', False))

        # 11. AFFIDABILITÀ E RESPONSABILITÀ
        self._replace_text(doc, '{{affidabilita}}', us_data.get('affidabilita_stratigrafica', ''))
        self._replace_text(doc, '{{responsabile_scientifico}}', us_data.get('responsabile_scientifico', ''))
        self._replace_text(doc, '{{data_rilevamento}}', self._format_date(us_data.get('data_rilevamento')))
        self._replace_text(doc, '{{responsabile_compilazione}}', us_data.get('responsabile_compilazione', ''))
        self._replace_text(doc, '{{data_rielaborazione}}', self._format_date(us_data.get('data_rielaborazione')))
        self._replace_text(doc, '{{responsabile_rielaborazione}}', us_data.get('responsabile_rielaborazione', ''))

        # Salva documento compilato
        doc.save(output_path)
        return output_path

    def export_usm(self, usm_data: Dict[str, Any], output_path: str) -> str:
        """
        Compila template USM con dati database
        Stessa logica di export_us ma per USM
        """
        doc = Document(self.template_path)

        # COMPILA CAMPI USM (simile a US ma con campi specifici USM)

        # Intestazione
        self._replace_text(doc, '{{usm_code}}', usm_data.get('usm_code', ''))
        self._replace_text(doc, '{{ente_responsabile}}', usm_data.get('ente_responsabile', ''))
        self._replace_text(doc, '{{anno}}', str(usm_data.get('anno', '')))

        # Localizzazione (uguale a US)
        self._replace_text(doc, '{{localita}}', usm_data.get('localita', ''))
        self._replace_text(doc, '{{area_struttura}}', usm_data.get('area_struttura', ''))

        # Documentazione
        self._replace_text(doc, '{{piante}}', usm_data.get('piante_riferimenti', ''))
        self._replace_text(doc, '{{prospetti}}', usm_data.get('prospetti_riferimenti', ''))
        self._replace_text(doc, '{{sezioni}}', usm_data.get('sezioni_riferimenti', ''))

        # CAMPI SPECIFICI USM
        self._replace_text(doc, '{{misure}}', usm_data.get('misure', ''))
        self._replace_text(doc, '{{superficie}}', str(usm_data.get('superficie_analizzata', '')))
        self._replace_text(doc, '{{definizione}}', usm_data.get('definizione', ''))

        # Tecnica costruttiva
        self._replace_text(doc, '{{tecnica}}', usm_data.get('tecnica_costruttiva', ''))
        self._replace_text(doc, '{{sezione_tipo}}', usm_data.get('sezione_muraria_tipo', ''))
        self._replace_text(doc, '{{sezione_spessore}}', usm_data.get('sezione_muraria_spessore', ''))
        self._replace_text(doc, '{{funzione_statica}}', usm_data.get('funzione_statica', ''))
        self._replace_text(doc, '{{modulo}}', usm_data.get('modulo', ''))

        # Materiali (JSON fields - format as text)
        laterizi = usm_data.get('materiali_laterizi', {})
        self._replace_text(doc, '{{laterizi}}', self._format_json_list(laterizi))

        litici = usm_data.get('materiali_elementi_litici', {})
        self._replace_text(doc, '{{litici}}', self._format_json_list(litici))

        legante = usm_data.get('legante', {})
        self._replace_text(doc, '{{legante}}', self._format_json_dict(legante))

        # Conservazione e altri campi
        self._replace_text(doc, '{{conservazione}}', usm_data.get('stato_conservazione', ''))
        self._replace_text(doc, '{{orientamento}}', usm_data.get('orientamento', ''))
        self._replace_text(doc, '{{uso_primario}}', usm_data.get('uso_primario', ''))

        # Sequenza fisica (uguale a US)
        self._compile_matrix_harris_table(doc, usm_data.get('sequenza_fisica', {}))

        # Descrizione
        self._replace_text(doc, '{{descrizione}}', usm_data.get('descrizione', ''))
        self._replace_text(doc, '{{interpretazione}}', usm_data.get('interpretazione', ''))

        # Datazione
        self._replace_text(doc, '{{datazione}}', usm_data.get('datazione', ''))
        self._replace_text(doc, '{{periodo}}', usm_data.get('periodo', ''))

        # Campionature USM
        campionature = usm_data.get('campionature', {})
        self._replace_checkbox(doc, '{{campione_litici}}', campionature.get('elementi_litici', False))
        self._replace_checkbox(doc, '{{campione_laterizi}}', campionature.get('laterizi', False))
        self._replace_checkbox(doc, '{{campione_malta}}', campionature.get('malta', False))

        # Responsabilità
        self._replace_text(doc, '{{responsabile_scientifico}}', usm_data.get('responsabile_scientifico', ''))
        self._replace_text(doc, '{{data_rilevamento}}', self._format_date(usm_data.get('data_rilevamento')))

        doc.save(output_path)
        return output_path

    # ===== METODI HELPER PER MANIPOLAZIONE DOCUMENTO =====

    def _replace_text(self, doc: Document, placeholder: str, value: str):
        """
        Sostituisce placeholder in TUTTO il documento (paragrafi, tabelle, header, footer)
        Mantiene formattazione originale
        """
        value = str(value) if value is not None else ''

        # Sostituisci in paragrafi
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                self._replace_in_paragraph(paragraph, placeholder, value)

        # Sostituisci in tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, value)

        # Sostituisci in header/footer
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                if placeholder in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder, value)

            # Footer
            for paragraph in section.footer.paragraphs:
                if placeholder in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder, value)

    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str):
        """
        Sostituisce testo nel paragrafo mantenendo TUTTA la formattazione originale
        """
        # IMPORTANTE: Sostituisci nel testo completo mantenendo i run
        if placeholder not in paragraph.text:
            return

        # Trova tutti i run che contengono parte del placeholder
        full_text = ''
        run_texts = []

        for run in paragraph.runs:
            run_texts.append(run.text)
            full_text += run.text

        # Se il placeholder è presente, sostituisci
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, value)

            # Cancella tutti i run esistenti tranne il primo
            for _ in range(len(paragraph.runs) - 1):
                paragraph._element.remove(paragraph.runs[-1]._element)

            # Aggiorna il primo run con il testo nuovo
            if paragraph.runs:
                paragraph.runs[0].text = new_text

    def _replace_checkbox(self, doc: Document, placeholder: str, checked: bool):
        """
        Sostituisce placeholder checkbox con simbolo ☑ o ☐
        """
        symbol = '☑' if checked else '☐'
        self._replace_text(doc, placeholder, symbol)

    def _compile_matrix_harris_table(self, doc: Document, sequenza_fisica: Dict):
        """
        Compila tabella Matrix Harris con le relazioni stratigrafiche
        """
        # Trova tabella Matrix Harris nel documento (cerca placeholder specifico)
        for table in doc.tables:
            # Cerca prima cella con "SEQUENZA FISICA" o simile
            if self._is_matrix_table(table):
                self._fill_matrix_table(table, sequenza_fisica)
                break

    def _is_matrix_table(self, table) -> bool:
        """Identifica se è la tabella Matrix Harris"""
        # Controlla se prima riga contiene intestazioni matrix
        if len(table.rows) > 0:
            first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
            return 'SEQUENZA' in first_row_text.upper() or 'MATRIX' in first_row_text.upper()
        return False

    def _fill_matrix_table(self, table, sequenza_fisica: Dict):
        """
        Compila celle tabella Matrix Harris con codici US/USM
        """
        # Mapping relazioni → righe tabella (dipende dal layout template)
        relations_map = {
            'uguale_a': 0,
            'si_lega_a': 1,
            'gli_si_appoggia': 2,
            'si_appoggia_a': 3,
            'coperto_da': 4,
            'copre': 5,
            'tagliato_da': 6,
            'taglia': 7,
            'riempito_da': 8,
            'riempie': 9
        }

        for relation, row_index in relations_map.items():
            if row_index < len(table.rows):
                values = sequenza_fisica.get(relation, [])
                # Scrivi nella seconda colonna (indice 1)
                if len(table.rows[row_index].cells) > 1:
                    cell = table.rows[row_index].cells[1]
                    cell.text = ', '.join(map(str, values)) if values else ''

    def _format_date(self, date_value) -> str:
        """Formatta data in formato italiano"""
        if not date_value:
            return ''

        if isinstance(date_value, str):
            return date_value

        try:
            return date_value.strftime('%d/%m/%Y')
        except:
            return str(date_value)

    def _format_json_list(self, json_data: Dict) -> str:
        """Formatta dizionario JSON come testo leggibile"""
        if not json_data:
            return ''

        lines = []
        for key, value in json_data.items():
            if isinstance(value, list):
                lines.append(f"{key}: {', '.join(map(str, value))}")
            else:
                lines.append(f"{key}: {value}")

        return '\n'.join(lines)

    def _format_json_dict(self, json_data: Dict) -> str:
        """Formatta dizionario JSON come testo"""
        if not json_data:
            return ''

        return '; '.join([f"{k}: {v}" for k, v in json_data.items()])


# ===== FUNZIONE HELPER PER FASTAPI =====

def create_us_word_from_template(us_instance, template_docx_path: str, output_dir: str) -> str:
    """
    Crea documento Word compilato da istanza US database

    Args:
        us_instance: Istanza UnitaStratigrafica dal database
        template_docx_path: Path al template .docx originale
        output_dir: Directory dove salvare output

    Returns:
        Path del file generato
    """
    exporter = USWordExporter(template_docx_path)

    # Converti istanza database in dizionario
    us_data = {
        'us_code': us_instance.us_code,
        'ente_responsabile': us_instance.ente_responsabile,
        'anno': us_instance.anno,
        'identificativo_rif': us_instance.identificativo_rif,
        'localita': us_instance.localita,
        'area_struttura': us_instance.area_struttura,
        'ambiente_unita_funzione': us_instance.ambiente_unita_funzione,
        'settori': us_instance.settori,
        'piante_riferimenti': us_instance.piante_riferimenti,
        'prospetti_riferimenti': us_instance.prospetti_riferimenti,
        'sezioni_riferimenti': us_instance.sezioni_riferimenti,
        'definizione': us_instance.definizione,
        'criteri_distinzione': us_instance.criteri_distinzione,
        'modo_formazione': us_instance.modo_formazione,
        'componenti_inorganici': us_instance.componenti_inorganici,
        'componenti_organici': us_instance.componenti_organici,
        'consistenza': us_instance.consistenza,
        'colore': us_instance.colore,
        'misure': us_instance.misure,
        'stato_conservazione': us_instance.stato_conservazione,
        'sequenza_fisica': us_instance.sequenza_fisica,
        'descrizione': us_instance.descrizione,
        'osservazioni': us_instance.osservazioni,
        'interpretazione': us_instance.interpretazione,
        'datazione': us_instance.datazione,
        'periodo': us_instance.periodo,
        'fase': us_instance.fase,
        'elementi_datanti': us_instance.elementi_datanti,
        'dati_quantitativi_reperti': us_instance.dati_quantitativi_reperti,
        'campionature': us_instance.campionature,
        'affidabilita_stratigrafica': us_instance.affidabilita_stratigrafica,
        'responsabile_scientifico': us_instance.responsabile_scientifico,
        'data_rilevamento': us_instance.data_rilevamento,
        'responsabile_compilazione': us_instance.responsabile_compilazione,
        'data_rielaborazione': us_instance.data_rielaborazione,
        'responsabile_rielaborazione': us_instance.responsabile_rielaborazione
    }

    # Genera nome file output
    filename = f"US_{us_instance.us_code}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(output_dir, filename)

    # Compila template
    return exporter.export_us(us_data, output_path)
