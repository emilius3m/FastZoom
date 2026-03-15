"""
app/services/tma_export_service.py

Servizio export Schede TMA ICCD 3.00
- PDF (reportlab) conforme al layout ministeriale ICCD
- Word (.docx, python-docx) conforme al layout ministeriale ICCD
"""

import io
from datetime import datetime
from typing import Dict, Any, List, Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, Image as RLImage,
)
from reportlab.pdfgen import canvas
from reportlab.platypus.flowables import HRFlowable
from loguru import logger

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont('Helvetica', 9)
        self.drawRightString(21*cm - 2.5*cm, 29.7*cm - 1.5*cm, f"Pagina {self._pageNumber} di {page_count}")
        self.restoreState()


try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx non disponibile per export Word TMA")


# ═══════════════════════════════════════════════
#  COLORI CONDIVISI
# ═══════════════════════════════════════════════
_PDF_COLORS = {
    'header_bg': colors.HexColor('#1a3a52'),
    'accent': colors.HexColor('#2c5aa0'),
    'accent_light': colors.HexColor('#e8eef5'),
    'border': colors.HexColor('#4a7ba7'),
    'text': colors.HexColor('#1a1a1a'),
    'grey': colors.HexColor('#666666'),
    'light_grey': colors.HexColor('#f5f5f5'),
    'section_bg': colors.HexColor('#edf2f7'),
}

_WORD_COLORS = {
    'header_bg': RGBColor(26, 58, 82),
    'accent': RGBColor(44, 90, 160),
    'accent_light': RGBColor(232, 238, 245),
    'text': RGBColor(26, 26, 26),
    'grey': RGBColor(100, 100, 100),
}


# ═══════════════════════════════════════════════
#  HELPER – sezioni ICCD dal dict
# ═══════════════════════════════════════════════

def _iccd_sections(s: Dict[str, Any]) -> List[dict]:
    """
    Restituisce la lista ordinata di sezioni ICCD,
    ciascuna con titolo + lista di (etichetta, valore).
    """
    sections: List[dict] = []

    # --- CD - CODICI ---
    sections.append({
        'title': 'CD - CODICI',
        'fields': [
            ('TSK - Tipo Scheda', s.get('tsk', 'TMA')),
            ('LIR - Livello ricerca', s.get('lir', 'I')),
        ],
        'subsections': [{
            'title': 'NCT - CODICE UNIVOCO',
            'fields': [
                ('NCTR - Codice regione', s.get('nctr', '')),
                ('NCTN - Numero catalogo generale', s.get('nctn', '')),
            ],
        }],
        'extra_fields': [
            ('ESC - Ente schedatore', s.get('esc', '')),
            ('ECP - Ente competente', s.get('ecp', '')),
        ],
    })

    # --- OG - OGGETTO ---
    sections.append({
        'title': 'OG - OGGETTO',
        'subsections': [{
            'title': 'OGT - OGGETTO',
            'fields': [
                ('OGTD - Definizione', s.get('ogtd', '')),
                ('OGTM - Definizione materiale componente', s.get('ogtm', '')),
            ],
        }],
    })

    # --- LC - LOCALIZZAZIONE ---
    sections.append({
        'title': 'LC - LOCALIZZAZIONE GEOGRAFICO-AMMINISTRATIVA',
        'subsections': [{
            'title': 'PVC - LOCALIZZAZIONE GEOGRAFICO-AMMINISTRATIVA ATTUALE',
            'fields': [
                ('PVCS - Stato', s.get('pvcs', 'ITALIA')),
                ('PVCR - Regione', s.get('pvcr', '')),
                ('PVCP - Provincia', s.get('pvcp', '')),
                ('PVCC - Comune', s.get('pvcc', '')),
            ],
        }],
    })

    # --- LDC - COLLOCAZIONE SPECIFICA ---
    ldc_fields = []
    if s.get('ldct'): ldc_fields.append(('LDCT - Tipologia', s['ldct']))
    if s.get('ldcn'): ldc_fields.append(('LDCN - Denominazione attuale', s['ldcn']))
    if s.get('ldcu'): ldc_fields.append(('LDCU - Indirizzo', s['ldcu']))
    if s.get('ldcs'): ldc_fields.append(('LDCS - Specifiche e note', s['ldcs']))
    if ldc_fields:
        sections.append({
            'title': 'LDC - COLLOCAZIONE SPECIFICA',
            'fields': ldc_fields,
        })

    # --- LA - ALTRE LOCALIZZAZIONI ---
    altre_loc = s.get('altre_localizzazioni') or []
    for loc in altre_loc:
        la_fields = []
        if loc.get('tcl'):  la_fields.append(('TCL - Tipo di localizzazione', loc['tcl']))
        la_sub_fields = []
        for k, lbl in [('prvs', 'PRVS - Stato'), ('prvr', 'PRVR - Regione'),
                        ('prvp', 'PRVP - Provincia'), ('prvc', 'PRVC - Comune')]:
            v = loc.get(k)
            if v: la_sub_fields.append((lbl, v))
        prc_fields = []
        if loc.get('prcu'): prc_fields.append(('PRCU - Denominazione', loc['prcu']))

        sec = {'title': 'LA - ALTRE LOCALIZZAZIONI GEOGRAFICO-AMMINISTRATIVE', 'fields': la_fields}
        subs = []
        if la_sub_fields:
            subs.append({'title': 'PRV - LOCALIZZAZIONE GEOGRAFICO-AMMINISTRATIVA', 'fields': la_sub_fields})
        if prc_fields:
            subs.append({'title': 'PRC - COLLOCAZIONE SPECIFICA', 'fields': prc_fields})
        if subs:
            sec['subsections'] = subs
        sections.append(sec)

    # --- RE / DSC - DATI DI SCAVO ---
    dsc_fields = []
    for k, lbl in [('scan', 'SCAN - Denominazione dello scavo'),
                    ('dscf', 'DSCF - Ente responsabile'),
                    ('dsca', 'DSCA - Responsabile scientifico'),
                    ('dsct', 'DSCT - Motivo'),
                    ('dscm', 'DSCM - Metodo'),
                    ('dscd', 'DSCD - Data'),
                    ('dscu', 'DSCU - Unità Stratigrafica'),
                    ('dscn', 'DSCN - Specifiche')]:
        v = s.get(k)
        if v: dsc_fields.append((lbl, v))
    if dsc_fields:
        sections.append({
            'title': "RE - MODALITA' DI REPERIMENTO",
            'subsections': [{
                'title': 'DSC - DATI DI SCAVO',
                'fields': dsc_fields,
            }],
        })

    # --- DT - CRONOLOGIA ---
    dt_fields = [('DTZG - Fascia cronologica di riferimento', s.get('dtzg', ''))]
    dtm = s.get('dtm') or []
    if dtm:
        dt_fields.append(('DTM - Motivazione cronologia', ', '.join(dtm)))
    sections.append({
        'title': 'DT - CRONOLOGIA',
        'subsections': [{
            'title': 'DTZ - CRONOLOGIA GENERICA',
            'fields': dt_fields,
        }],
    })

    # --- DA - DATI ANALITICI ---
    if s.get('nsc'):
        sections.append({
            'title': 'DA - DATI ANALITICI',
            'fields': [('NSC - Notizie storico-critiche', s['nsc'])],
        })

    # --- MA - MATERIALE (repeating) ---
    materiali = s.get('materiali') or []
    for mat in materiali:
        ma_fields = []
        for k, lbl in [('macc', 'MACC - Categoria'), ('macl', 'MACL - Classe'),
                        ('macd', 'MACD - Definizione'), ('macp', 'MACP - Precisazione tipologica'),
                        ('macq', 'MACQ - Quantità'), ('mas', 'MAS - Specifiche')]:
            v = mat.get(k)
            if v is not None and str(v).strip():
                ma_fields.append((lbl, str(v)))
        sections.append({
            'title': 'MA - MATERIALE',
            'subsections': [{
                'title': 'MAC - MATERIALE COMPONENTE',
                'fields': ma_fields,
            }],
        })

    # --- TU - CONDIZIONE GIURIDICA ---
    sections.append({
        'title': 'TU - CONDIZIONE GIURIDICA E VINCOLI',
        'subsections': [{
            'title': 'CDG - CONDIZIONE GIURIDICA',
            'fields': [('CDGG - Indicazione generica', s.get('cdgg', ''))],
        }],
    })

    # --- DO - FONTI E DOCUMENTI ---
    fotografie = s.get('fotografie') or []
    if fotografie:
        fta_subs = []
        for foto in fotografie:
            fta_fields = []
            if foto.get('ftax'): fta_fields.append(('FTAX - Genere', foto['ftax']))
            if foto.get('ftap'): fta_fields.append(('FTAP - Tipo', foto['ftap']))
            if foto.get('ftan'): fta_fields.append(('FTAN - Codice identificativo', foto['ftan']))
            if fta_fields:
                fta_subs.append({'title': 'FTA - DOCUMENTAZIONE FOTOGRAFICA', 'fields': fta_fields})
        if fta_subs:
            sections.append({
                'title': 'DO - FONTI E DOCUMENTI DI RIFERIMENTO',
                'subsections': fta_subs,
            })

    # --- AD - ACCESSO AI DATI ---
    ad_fields = [('ADSP - Profilo di accesso', str(s.get('adsp', 2)))]
    if s.get('adsm'):
        ad_fields.append(('ADSM - Motivazione', s['adsm']))
    sections.append({
        'title': 'AD - ACCESSO AI DATI',
        'subsections': [{
            'title': 'ADS - SPECIFICHE DI ACCESSO AI DATI',
            'fields': ad_fields,
        }],
    })

    # --- CM - COMPILAZIONE ---
    cm_fields = [('CMPD - Data', s.get('cmpd', ''))]
    cmpn = s.get('cmpn') or []
    for nome in cmpn:
        cm_fields.append(('CMPN - Nome', nome))
    fur = s.get('fur') or []
    for nome in fur:
        cm_fields.append(('FUR - Funzionario responsabile', nome))
    sections.append({
        'title': 'CM - COMPILAZIONE',
        'subsections': [{
            'title': 'CMP - COMPILAZIONE',
            'fields': cm_fields,
        }],
    })

    return sections


def _collect_export_photos(scheda: Dict[str, Any]) -> List[Dict[str, Any]]:
    photos: List[Dict[str, Any]] = []
    for idx, foto in enumerate((scheda.get('fotografie') or []), start=1):
        image_bytes = foto.get('image_bytes')
        if isinstance(image_bytes, bytearray):
            image_bytes = bytes(image_bytes)

        if not isinstance(image_bytes, (bytes, bytearray)) or len(image_bytes) == 0:
            continue

        caption = foto.get('ftan') or foto.get('ftap') or foto.get('ftax') or f"Foto {idx}"
        photos.append({
            'index': idx,
            'caption': str(caption),
            'file_path': str(foto.get('file_path') or ''),
            'image_bytes': bytes(image_bytes),
        })
    return photos


def _add_pdf_photo_gallery(story: List[Any], scheda: Dict[str, Any], styles):
    photos = _collect_export_photos(scheda)
    if not photos:
        return

    story.append(PageBreak())
    story.append(Paragraph('DO - ALLEGATI FOTOGRAFICI', styles['TmaSectionHeading']))
    story.append(Spacer(1, 0.2 * cm))

    col_width = 8.2 * cm
    cells: List[Any] = []

    for photo in photos:
        try:
            img_stream = io.BytesIO(photo['image_bytes'])
            img = RLImage(img_stream)
            img._restrictSize(col_width - (0.5 * cm), 5.5 * cm)

            block = [
                Paragraph(f"<b>Foto {photo['index']}</b> — {photo['caption']}", styles['TmaLabel']),
                Spacer(1, 1.2 * mm),
                img,
            ]

            if photo['file_path']:
                block.extend([
                    Spacer(1, 1 * mm),
                    Paragraph(photo['file_path'], styles['TmaPageNum']),
                ])

            cells.append(block)
        except Exception:
            cells.append([
                Paragraph(f"<b>Foto {photo['index']}</b> — {photo['caption']}", styles['TmaLabel']),
                Paragraph("Immagine non disponibile per il rendering", styles['TmaValue']),
            ])

    rows = []
    for i in range(0, len(cells), 2):
        row = cells[i:i + 2]
        if len(row) < 2:
            row.append("")
        rows.append(row)

    gallery = Table(rows, colWidths=[col_width, col_width], hAlign='LEFT')
    gallery.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('BOX', (0, 0), (-1, -1), 0.5, _PDF_COLORS['border']),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, _PDF_COLORS['light_grey']),
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
    ]))

    story.append(gallery)
    story.append(Spacer(1, 0.2 * cm))


# ═══════════════════════════════════════════════
#  PDF EXPORT (reportlab)
# ═══════════════════════════════════════════════

def _setup_pdf_styles():
    styles = getSampleStyleSheet()

    def add_or_update(name, **kw):
        if name in styles.byName:
            for k, v in kw.items():
                setattr(styles[name], k, v)
        else:
            styles.add(ParagraphStyle(name=name, **kw))

    add_or_update('TmaTitle', parent=styles['Heading1'],
                  fontSize=16, textColor=_PDF_COLORS['header_bg'],
                  alignment=TA_CENTER, fontName='Helvetica-Bold',
                  spaceAfter=4, spaceBefore=4)

    add_or_update('TmaSubtitle', parent=styles['Normal'],
                  fontSize=10, textColor=_PDF_COLORS['grey'],
                  alignment=TA_CENTER, fontName='Helvetica-Oblique',
                  spaceAfter=8)

    add_or_update('TmaSectionHeading', parent=styles['Normal'],
                  fontSize=10, textColor=colors.white,
                  fontName='Helvetica-Bold', spaceAfter=2, spaceBefore=8,
                  backColor=_PDF_COLORS['accent'], leftIndent=4,
                  rightIndent=4, leading=14)

    add_or_update('TmaSubsection', parent=styles['Normal'],
                  fontSize=9, textColor=_PDF_COLORS['header_bg'],
                  fontName='Helvetica-Bold', spaceAfter=2, spaceBefore=4,
                  leftIndent=8, leading=12)

    add_or_update('TmaLabel', parent=styles['Normal'],
                  fontSize=8, fontName='Helvetica-Bold',
                  textColor=_PDF_COLORS['text'], spaceAfter=2, leading=10)

    add_or_update('TmaValue', parent=styles['Normal'],
                  fontSize=9, fontName='Helvetica',
                  textColor=_PDF_COLORS['text'], spaceAfter=6,
                  alignment=TA_LEFT, leading=13)

    add_or_update('TmaPageNum', parent=styles['Normal'],
                  fontSize=7, textColor=_PDF_COLORS['grey'],
                  alignment=TA_CENTER, fontName='Helvetica-Oblique')

    add_or_update('TmaFooter', parent=styles['Normal'],
                  fontSize=7, textColor=_PDF_COLORS['grey'],
                  alignment=TA_CENTER, fontName='Helvetica-Oblique',
                  spaceBefore=12)

    return styles


def generate_tma_pdf(scheda: Dict[str, Any]) -> bytes:
    """
    Genera PDF conforme ICCD per una scheda TMA.
    Riceve il dict serializzato (SchedaTMARead-like).
    """
    styles = _setup_pdf_styles()
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        title=f"Scheda TMA — NCT {scheda.get('nctr', '')}{scheda.get('nctn', '')}",
    )

    story: list = []

    # Titolo
    story.append(Paragraph("Scheda", styles['TmaTitle']))
    story.append(HRFlowable(width="100%", thickness=1.5, color=_PDF_COLORS['accent']))
    story.append(Spacer(1, 0.4 * cm))

    def _make_fields_table(fields_list):
        if not fields_list: return None
        cells = []
        for lbl, val in fields_list:
            cells.append([
                Paragraph(lbl, styles['TmaLabel']),
                Paragraph(str(val) if val else "-", styles['TmaValue'])
            ])
        rows = []
        for i in range(0, len(cells), 2):
            row = cells[i:i+2]
            if len(row) == 1:
                row.append([])
            rows.append(row)
        if rows:
            cw = (doc.width) / 2.0
            t = Table(rows, colWidths=[cw, cw])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
            ]))
            return t
        return None

    # Sezioni ICCD
    for sec in _iccd_sections(scheda):
        if sec['title'] == 'MA - MATERIALE':
            story.append(HRFlowable(width="100%", thickness=0.5, color=_PDF_COLORS['light_grey']))
            story.append(Spacer(1, 0.4 * cm))

        # Intestazione sezione
        story.append(Paragraph(sec['title'], styles['TmaSectionHeading']))

        # Campi diretti della sezione
        t = _make_fields_table(sec.get('fields', []))
        if t: story.append(t)

        # Extra fields (usati per ESC/ECP dopo NCT)
        t_ext = _make_fields_table(sec.get('extra_fields', []))
        if t_ext: story.append(t_ext)

        # Sottosezioni
        for sub in sec.get('subsections', []):
            story.append(Paragraph(sub['title'], styles['TmaSubsection']))
            t_sub = _make_fields_table(sub.get('fields', []))
            if t_sub: story.append(t_sub)

        story.append(Spacer(1, 0.6 * cm))

    # Allegati fotografici (se presenti bytes immagine)
    _add_pdf_photo_gallery(story, scheda, styles)

    # Entita Multimediali associate (metadati)
    photos = _collect_export_photos(scheda)
    if photos:
        story.append(PageBreak())
        story.append(Paragraph("Entita' multimediali associate", styles['TmaSectionHeading']))
        story.append(Spacer(1, 0.4 * cm))
        for photo in photos:
            story.append(Paragraph("MC - METADATI DI CATALOGAZIONE", styles['TmaSubsection']))
            story.append(Paragraph("FTA - DOCUMENTAZIONE FOTOGRAFICA", styles['TmaSubsection']))
            
            fta_fields = []
            fta_fields.append(('FTAP - Tipo', 'fotografia digitale (file)'))
            fta_fields.append(('FTAN - Codice identificativo', photo['caption']))
            t_fta = _make_fields_table(fta_fields)
            if t_fta: story.append(t_fta)
            
            story.append(Paragraph("MM - METADATI DATO MULTIMEDIALE", styles['TmaSubsection']))
            story.append(Paragraph("MMT - METADATI TECNICI DATO MULTIMEDIALE", styles['TmaSubsection']))
            
            mm_fields = []
            filename = photo['file_path'].split('/')[-1] if photo['file_path'] else f"{photo['caption']}.jpg"
            mm_fields.append(('MMTO - Nome file originale', filename))
            t_mm = _make_fields_table(mm_fields)
            if t_mm: story.append(t_mm)
            
            story.append(Spacer(1, 0.4 * cm))
            story.append(HRFlowable(width="50%", thickness=0.5, color=_PDF_COLORS['light_grey']))
            story.append(Spacer(1, 0.6 * cm))

    # Footer / Firma
    story.append(PageBreak())
    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_PDF_COLORS['border']))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Firma", styles['TmaLabel']))
    story.append(Spacer(1, 1.5 * cm))

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    logger.info(f"✓ TMA PDF generato: NCT {scheda.get('nctr','')}{scheda.get('nctn','')} ({len(pdf_bytes)} bytes)")
    return pdf_bytes


# ═══════════════════════════════════════════════
#  WORD EXPORT (python-docx)
# ═══════════════════════════════════════════════

def _rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_word_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def _set_word_cell_border(cell, color: str = "B7C7D9", size: int = 4):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    for edge in ('top', 'left', 'bottom', 'right'):
        tag = f'w:{edge}'
        elem = tc_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tc_borders.append(elem)
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)


def _add_word_title_block(doc: Document, scheda: Dict[str, Any]):
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    _set_word_cell_shading(title_cell, _rgb_to_hex(_WORD_COLORS['header_bg']))
    _set_word_cell_border(title_cell, color="1A3A52", size=8)

    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SCHEDA TMA")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)

    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Tabella Materiali Archeologici — ICCD 3.00")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(230, 240, 250)

    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = "Table Grid"

    nct = f"{scheda.get('nctr', '')}{scheda.get('nctn', '')}"
    meta_left = meta_table.cell(0, 0)
    meta_right = meta_table.cell(0, 1)

    _set_word_cell_shading(meta_left, _rgb_to_hex(_WORD_COLORS['accent_light']))
    _set_word_cell_shading(meta_right, _rgb_to_hex(_WORD_COLORS['accent_light']))

    pl = meta_left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = pl.add_run(f"NCT: {nct or '-'}")
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = _WORD_COLORS['header_bg']

    pr = meta_right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = pr.add_run(f"Export: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r2.font.size = Pt(8)
    r2.font.color.rgb = _WORD_COLORS['grey']

    doc.add_paragraph()


def _add_word_section_table(doc: Document, section: Dict[str, Any]):
    # Heading ribbon
    heading_table = doc.add_table(rows=1, cols=1)
    heading_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    heading_cell = heading_table.cell(0, 0)
    _set_word_cell_shading(heading_cell, _rgb_to_hex(_WORD_COLORS['accent']))
    _set_word_cell_border(heading_cell, color="2C5AA0", size=8)

    hp = heading_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(section['title'])
    hr.font.size = Pt(10)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(255, 255, 255)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    def add_subsection_row(title: str):
        row_cells = table.add_row().cells
        merged = row_cells[0].merge(row_cells[1])
        _set_word_cell_shading(merged, _rgb_to_hex(_WORD_COLORS['accent_light']))
        _set_word_cell_border(merged, color="9EB6CF", size=4)
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = _WORD_COLORS['header_bg']

    def add_field_row(label: str, value: Any):
        row_cells = table.add_row().cells
        label_cell, value_cell = row_cells[0], row_cells[1]

        _set_word_cell_shading(label_cell, "F5F8FC")
        _set_word_cell_border(label_cell, color="C5D3E3", size=4)
        _set_word_cell_border(value_cell, color="C5D3E3", size=4)

        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(str(label))
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = _WORD_COLORS['header_bg']

        vp = value_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(str(value) if value not in (None, "") else "-")
        vr.font.size = Pt(9)
        vr.font.color.rgb = _WORD_COLORS['text']

    for label, value in section.get('fields', []):
        add_field_row(label, value)

    for label, value in section.get('extra_fields', []):
        add_field_row(label, value)

    for sub in section.get('subsections', []):
        add_subsection_row(sub['title'])
        for label, value in sub.get('fields', []):
            add_field_row(label, value)

    doc.add_paragraph()


def _add_word_photo_gallery(doc: Document, scheda: Dict[str, Any]):
    photos = _collect_export_photos(scheda)
    if not photos:
        return

    heading_table = doc.add_table(rows=1, cols=1)
    heading_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    heading_cell = heading_table.cell(0, 0)
    _set_word_cell_shading(heading_cell, _rgb_to_hex(_WORD_COLORS['accent']))
    _set_word_cell_border(heading_cell, color="2C5AA0", size=8)

    hp = heading_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("DO - ALLEGATI FOTOGRAFICI")
    hr.font.size = Pt(10)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    for photo in photos:
        cap = doc.add_paragraph()
        cap_run = cap.add_run(f"Foto {photo['index']} — {photo['caption']}")
        cap_run.font.size = Pt(9)
        cap_run.font.bold = True
        cap_run.font.color.rgb = _WORD_COLORS['header_bg']

        try:
            doc.add_picture(io.BytesIO(photo['image_bytes']), width=Cm(8.5))
        except Exception:
            p = doc.add_paragraph("Immagine non disponibile per il rendering")
            p_run = p.runs[0]
            p_run.font.size = Pt(8)
            p_run.font.italic = True
            p_run.font.color.rgb = _WORD_COLORS['grey']

        if photo['file_path']:
            meta = doc.add_paragraph(photo['file_path'])
            meta_run = meta.runs[0]
            meta_run.font.size = Pt(7)
            meta_run.font.color.rgb = _WORD_COLORS['grey']

        doc.add_paragraph()


def generate_tma_word(scheda: Dict[str, Any]) -> bytes:
    """
    Genera documento Word (.docx) conforme ICCD per una scheda TMA.
    Riceve il dict serializzato (SchedaTMARead-like).
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx non installato")

    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _add_word_title_block(doc, scheda)

    # Sezioni ICCD
    for sec in _iccd_sections(scheda):
        _add_word_section_table(doc, sec)

    # Allegati fotografici (se presenti bytes immagine)
    _add_word_photo_gallery(doc, scheda)

    # Firma
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma_run = firma.add_run("Firma")
    firma_run.font.size = Pt(9)
    firma_run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"Documento generato da FastZoom Archaeological System — {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    footer_run.font.size = Pt(7)
    footer_run.font.italic = True
    footer_run.font.color.rgb = _WORD_COLORS['grey']

    # Salva
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    word_bytes = buf.getvalue()
    buf.close()
    logger.info(f"✓ TMA Word generato: NCT {scheda.get('nctr','')}{scheda.get('nctn','')} ({len(word_bytes)} bytes)")
    return word_bytes
