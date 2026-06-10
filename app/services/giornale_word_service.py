"""
app/services/giornale_word_service_v2.py

Servizio Word per Giornale di Cantiere - VERSIONE 2.0 COMPLETA
- Design professionale con formattazione avanzata
- Tutte le 11 sezioni documentate
- 100% modificabile
- Zero perdita di dati
- Conforme agli standard ICCD

Autore: FastZoom Archaeological System
Data: 13 Novembre 2025
"""

import io
from datetime import datetime, date
from typing import Dict, Any, List, Optional
from loguru import logger

from app.services.giornale_export_builder import build_giornale_export_model

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - install with: pip install python-docx")
    DOCX_AVAILABLE = False


class GiornaleWordGeneratorV2:
    """Generatore Word professionali per Giornale di Cantiere - Versione 2.0"""

    # Colori professionali
    COLOR_HEADER_BG = RGBColor(26, 58, 82)
    COLOR_HEADER_TEXT = RGBColor(255, 255, 255)
    COLOR_ACCENT = RGBColor(44, 90, 160)
    COLOR_TEXT = RGBColor(26, 26, 26)
    COLOR_GREY = RGBColor(100, 100, 100)

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required: pip install python-docx")

    def generate_giornale_word(self,
                              giornali: List[Dict[str, Any]],
                              cantiere_info: Dict[str, Any],
                              site_info: Dict[str, Any]) -> bytes:
        """
        Genera documento Word completo del giornale di cantiere
        
        Include tutte le 11 sezioni:
        1. Intestazione progetto
        2. Informazioni generali
        3. Condizioni meteorologiche
        4. Descrizione lavori
        5. Risorse impiegate
        6. Unità stratigrafiche
        7. Materiali rinvenuti
        8. Documentazione
        9. Disposizioni
        10. Eventi particolari
        11. Note e validazione
        """
        try:
            model = build_giornale_export_model(giornali, cantiere_info, site_info)
            doc = Document()
            
            # Margini professionali
            for section in doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(1.6)
                section.right_margin = Cm(1.6)

            self._render_giornale_export_word(doc, model)

            # Salva in buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            word_bytes = buffer.getvalue()
            buffer.close()

            logger.info(f"✓ Word generato: {cantiere_info.get('nome')} ({len(word_bytes)} bytes)")
            return word_bytes

        except Exception as e:
            logger.error(f"✗ Errore generazione Word: {e}")
            raise

    def _render_giornale_export_word(self, doc, model: Dict[str, Any]) -> None:
        self._word_title(doc, model["title"], size=18)
        self._word_subtitle(doc, model["subtitle"])
        self._word_divider(doc)
        self._word_key_value_table(doc, model["cover_rows"])

        doc.add_page_break()
        self._word_title(doc, "SCHEDA CANTIERE", size=16)
        for section in model["summary_sections"]:
            self._word_add_section(doc, section)

        for entry in model["giornali"]:
            doc.add_page_break()
            self._word_title(doc, entry["heading"].upper(), size=15)
            for section in entry["sections"]:
                self._word_add_section(doc, section)

        doc.add_page_break()
        self._word_title(doc, "FIRME E VALIDAZIONI", size=16)
        p = doc.add_paragraph("Le firme sotto riportate attestano la presa visione del registro esportato.")
        self._set_paragraph_font(p, 9)
        signature_rows = [
            (label, "Firma ______________________________    Data __________")
            for label, _ in model["signature_rows"]
        ]
        self._word_key_value_table(doc, signature_rows)
        doc.add_paragraph()
        footer = doc.add_paragraph(
            f"Documento generato da FastZoom il {model['generated_at'].strftime('%d/%m/%Y alle %H:%M')}"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_font(footer, 8, italic=True, color=self.COLOR_GREY)

    def _word_add_section(self, doc, section: Dict[str, Any]) -> None:
        self._add_section_heading(doc, section["title"])
        kind = section.get("kind")

        if kind == "table":
            self._word_key_value_table(doc, section.get("rows", []))
        elif kind == "text":
            items = section.get("items", [])
            if items:
                for label, text in items:
                    p = doc.add_paragraph()
                    label_run = p.add_run(f"{label}: ")
                    label_run.bold = True
                    p.add_run(str(text))
                    self._set_paragraph_font(p, 9)
            else:
                self._set_paragraph_font(doc.add_paragraph("N/D"), 9)
        elif kind == "mixed":
            for label, text in section.get("items", []):
                p = doc.add_paragraph()
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                p.add_run(str(text))
                self._set_paragraph_font(p, 9)
            for table in section.get("tables", []):
                p = doc.add_paragraph()
                run = p.add_run(table["title"])
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = self.COLOR_ACCENT
                self._word_plain_table(doc, table.get("headers", []), table.get("rows", []))
            if section.get("photos"):
                p = doc.add_paragraph()
                run = p.add_run("Immagini")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = self.COLOR_ACCENT
                self._word_add_photos(doc, section["photos"])
        doc.add_paragraph()

    def _word_title(self, doc, text: str, size: int = 18) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_HEADER_BG

    def _word_subtitle(self, doc, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = self.COLOR_GREY

    def _word_divider(self, doc) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(17.8)
        cell = table.rows[0].cells[0]
        cell.text = ""
        self._set_cell_shading(cell, "2C5AA0")
        self._set_cell_borders(cell, "2C5AA0", size="8")
        cell.height = Cm(0.08)
        doc.add_paragraph()

    def _word_key_value_table(self, doc, rows) -> None:
        table = doc.add_table(rows=max(len(rows), 1), cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(5)
        table.columns[1].width = Cm(11)
        source_rows = rows or [("", "")]
        for idx, (label, value) in enumerate(source_rows):
            cells = table.rows[idx].cells
            cells[0].text = str(label)
            cells[1].text = str(value)
            self._format_cell(cells[0], bold=True, color=self.COLOR_HEADER_BG, shading="E8EEF5")
            self._format_cell(cells[1])
            self._set_cell_borders(cells[0])
            self._set_cell_borders(cells[1])

    def _word_plain_table(self, doc, headers: List[str], rows: List[List[str]]) -> None:
        headers = headers or [""]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = True
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = str(header)
            self._format_cell(table.rows[0].cells[idx], bold=True, color=self.COLOR_HEADER_TEXT, shading="2C5AA0")
            self._set_cell_borders(table.rows[0].cells[idx])
        for row_data in rows:
            row = table.add_row()
            for idx, value in enumerate(row_data[:len(headers)]):
                row.cells[idx].text = str(value)
                self._format_cell(row.cells[idx])
                self._set_cell_borders(row.cells[idx])

    def _word_add_photos(self, doc, photos: List[Dict[str, Any]]) -> None:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(8)
        table.columns[1].width = Cm(8)

        for idx in range(0, len(photos), 2):
            row = table.add_row()
            for cell_index in range(2):
                cell = row.cells[cell_index]
                self._set_cell_borders(cell)
                photo_index = idx + cell_index
                if photo_index >= len(photos):
                    cell.text = ""
                    continue
                self._word_fill_photo_cell(cell, photos[photo_index])

    def _word_fill_photo_cell(self, cell, photo: Dict[str, Any]) -> None:
        image_paragraph = cell.paragraphs[0]
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_bytes = photo.get("image_bytes")
        if image_bytes:
            try:
                image_paragraph.add_run().add_picture(io.BytesIO(image_bytes), width=Inches(2.75))
            except Exception as exc:
                logger.warning(f"Errore inserimento immagine Word: {exc}")
                image_paragraph.add_run("[Immagine non leggibile]")
        else:
            image_paragraph.add_run("[Immagine non disponibile nell'export]")

        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = photo.get("title") or "N/D"
        description = photo.get("description")
        caption_text = f"Foto {photo.get('index')}: {title}"
        if description and description != "N/D":
            caption_text += f"\n{description}"
        run = caption.add_run(caption_text)
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = self.COLOR_GREY

    def _word_add_photos_legacy(self, doc, photos: List[Dict[str, Any]]) -> None:
        for photo in photos:
            image_bytes = photo.get("image_bytes")
            if image_bytes:
                try:
                    doc.add_picture(io.BytesIO(image_bytes), width=Inches(5.2))
                except Exception as exc:
                    logger.warning(f"Errore inserimento immagine Word: {exc}")
                    self._set_paragraph_font(doc.add_paragraph("[Immagine non leggibile]"), 8, italic=True)
            else:
                self._set_paragraph_font(doc.add_paragraph("[Immagine non disponibile nell'export]"), 8, italic=True)

            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title = photo.get("title") or "N/D"
            description = photo.get("description")
            caption_text = f"Foto {photo.get('index')}: {title}"
            if description and description != "N/D":
                caption_text += f" - {description}"
            run = caption.add_run(caption_text)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = self.COLOR_GREY

    def _format_cell(self, cell, bold: bool = False, color=None, shading: Optional[str] = None) -> None:
        if shading:
            self._set_cell_shading(cell, shading)
        for paragraph in cell.paragraphs:
            self._set_paragraph_font(paragraph, 8, bold=bold, color=color)

    def _set_paragraph_font(self, paragraph, size: int, bold: bool = False, italic: bool = False, color=None) -> None:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

    def _set_cell_shading(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn('w:shd'))
        if shading is None:
            shading = OxmlElement('w:shd')
            tc_pr.append(shading)
        shading.set(qn('w:fill'), fill)

    def _set_cell_borders(self, cell, color: str = "4A7BA7", size: str = "4") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), size)
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), color)

    def _add_title_page(self, doc, cantiere_info, site_info, num_giornali):
        """Pagina titolo professionale con informazioni complete del cantiere"""
        
        # Titolo principale
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("GIORNALE DEI LAVORI DI CANTIERE")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = self.COLOR_HEADER_BG

        # Sottotitolo
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Documentazione Archeologica Conforme agli Standard ICCD")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = self.COLOR_GREY

        doc.add_paragraph()  # Spazio

        # Blocco informazioni stato e priorità
        status_table = doc.add_table(rows=2, cols=4)
        status_table.style = 'Light Grid Accent 1'
        
        status_data = [
            ("STATO CANTIERE:", cantiere_info.get('stato_formattato', 'N/D')),
            ("PRIORITÀ:", f"{cantiere_info.get('priorita', 'N/D')}/5"),
            ("DURATA:", f"{cantiere_info.get('durata_giorni', 'N/D')} giorni" if cantiere_info.get('durata_giorni') else "N/D"),
            ("CODICE:", cantiere_info.get('codice', 'N/D'))
        ]
        
        for i, (label, value) in enumerate(status_data):
            row_cells = status_table.rows[0 if i < 2 else 1].cells
            cell_idx = i if i < 2 else i - 2
            row_cells[cell_idx].text = f"{label} {value}"
            
            # Formattazione celle stato
            para = row_cells[cell_idx].paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(11)
                run.font.bold = True
                if i == 0:  # Stato - colore speciale
                    if cantiere_info.get('stato') == 'in_corso':
                        run.font.color.rgb = RGBColor(34, 197, 94)  # Verde
                    elif cantiere_info.get('stato') == 'sospeso':
                        run.font.color.rgb = RGBColor(251, 191, 36)  # Giallo
                    elif cantiere_info.get('stato') == 'completato':
                        run.font.color.rgb = RGBColor(107, 114, 128)  # Grigio
                    else:
                        run.font.color.rgb = RGBColor(59, 130, 246)  # Blu

        doc.add_paragraph()  # Spazio

        # Tabella informazioni principali ampliata
        header_table = doc.add_table(rows=10, cols=2)
        header_table.style = 'Light Grid Accent 1'

        header_rows = [
            ("OGGETTO:", cantiere_info.get('oggetto_appalto', cantiere_info.get('nome', 'N/D'))),
            ("COMMITTENTE:", cantiere_info.get('committente', 'N/D')),
            ("IMPRESA ESECUTRICE:", cantiere_info.get('impresa_esecutrice', 'N/D')),
            ("DIRETTORE DEI LAVORI:", cantiere_info.get('direttore_lavori', 'N/D')),
            ("RESPONSABILE PROCEDIMENTO:", cantiere_info.get('responsabile_procedimento', 'N/D')),
            ("RESPONSABILE CANTIERE:", cantiere_info.get('responsabile_cantiere', 'N/D')),
            ("TIPOLOGIA INTERVENTO:", cantiere_info.get('tipologia_intervento', 'N/D')),
            ("SITO ARCHEOLOGICO:", site_info.get('name', 'N/D')),
            ("DATA DOCUMENTO:", datetime.now().strftime('%d/%m/%Y %H:%M')),
            ("GIORNALI INCLUSI:", str(num_giornali)),
        ]

        for i, (label, value) in enumerate(header_rows):
            row_cells = header_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            
            # Formattazione
            label_para = row_cells[0].paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = self.COLOR_HEADER_BG
            
            value_para = row_cells[1].paragraphs[0]
            if value_para.runs:
                value_run = value_para.runs[0]
                value_run.font.size = Pt(10)

        doc.add_paragraph()
        
        # Sezione aggiuntiva - Informazioni geografiche e codici
        geo_table = doc.add_table(rows=2, cols=3)
        geo_table.style = 'Light Grid Accent 1'
        
        geo_data = [
            ("AREA:", cantiere_info.get('area_descrizione', 'N/D')),
            ("QUOTA:", cantiere_info.get('quota', 'N/D')),
            ("CODICE CUP:", cantiere_info.get('codice_cup', 'N/D')),
            ("CODICE CIG:", cantiere_info.get('codice_cig', 'N/D')),
            ("IMPORTO LAVORI:", f"€{cantiere_info.get('importo_lavori', 'N/D'):,.2f}" if cantiere_info.get('importo_lavori') else "N/D"),
            ("COORDINATE:", f"{cantiere_info.get('coordinate_lat', 'N/D')}, {cantiere_info.get('coordinate_lon', 'N/D')}" if cantiere_info.get('coordinate_lat') and cantiere_info.get('coordinate_lon') else "N/D")
        ]
        
        for i, (label, value) in enumerate(geo_data):
            row_cells = geo_table.rows[0 if i < 3 else 1].cells
            cell_idx = i if i < 3 else i - 3
            row_cells[cell_idx].text = f"{label} {value}"
            
            # Formattazione
            para = row_cells[cell_idx].paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.bold = True

        doc.add_paragraph()
        
        # Nota informativa
        nota = doc.add_paragraph()
        nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        nota_run = nota.add_run(
            "Questo documento contiene la documentazione completa delle attività di scavo "
            "conforme agli standard ICCD del Ministero della Cultura italiano. "
            "Tutti i dati sono tracciati e validabili. "
            f"Stato attuale cantiere: {cantiere_info.get('stato_formattato', 'N/D')}."
        )
        nota_run.font.size = Pt(9)
        nota_run.font.italic = True
        nota_run.font.color.rgb = self.COLOR_GREY

    def _add_index(self, doc, giornali):
        """Aggiunge indice dei giornali"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("INDICE")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = self.COLOR_HEADER_BG

        doc.add_paragraph()

        for i, g in enumerate(giornali, 1):
            data = self._format_date(g.get('data', 'N/D'))
            p = doc.add_paragraph(f"Giornale {i}: {data}", style='List Number')

    def _add_stato_cantiere_section(self, doc, cantiere_info):
        """Aggiunge sezione completa sullo stato del cantiere"""
        
        self._add_section_heading(doc, "STATO DEL CANTIERE E INFORMAZIONI CRITICHE")
        
        # Tabella stato e progressione
        status_table = doc.add_table(rows=6, cols=2)
        status_table.style = 'Light Grid Accent 1'
        
        # Dati per la tabella di stato
        status_rows = [
            ("STATO ATTUALE:", cantiere_info.get('stato_formattato', 'N/D')),
            ("PRIORITÀ INTERVENTO:", f"{cantiere_info.get('priorita', 'N/D')}/5" + " - ALTA" if cantiere_info.get('priorita', 0) >= 4 else " - MEDIA" if cantiere_info.get('priorita', 0) >= 2 else " - BASSA"),
            ("DURATA GIORNALIERA:", f"{cantiere_info.get('durata_giorni', 'N/D')} giorni" if cantiere_info.get('durata_giorni') else "In corso"),
            ("CANTIERE IN CORSO:", "SÌ" if cantiere_info.get('e_in_corso') else "NO"),
            ("CODICE IDENTIFICATIVO:", cantiere_info.get('codice', 'N/D')),
            ("RESPONSABILE CANTIERE:", cantiere_info.get('responsabile_cantiere', 'N/D'))
        ]
        
        for i, (label, value) in enumerate(status_rows):
            row_cells = status_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            
            # Formattazione etichette
            label_para = row_cells[0].paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = self.COLOR_ACCENT
            
            # Formattazione valori
            value_para = row_cells[1].paragraphs[0]
            if value_para.runs:
                value_run = value_para.runs[0]
                value_run.font.size = Pt(10)
                # Colore speciale per stato
                if i == 0:  # Stato
                    if cantiere_info.get('stato') == 'in_corso':
                        value_run.font.color.rgb = RGBColor(34, 197, 94)  # Verde
                    elif cantiere_info.get('stato') == 'sospeso':
                        value_run.font.color.rgb = RGBColor(251, 191, 36)  # Giallo
                    elif cantiere_info.get('stato') == 'completato':
                        value_run.font.color.rgb = RGBColor(107, 114, 128)  # Grigio

        doc.add_paragraph()

        # Tabella timeline - Programmato vs Effettivo
        self._add_section_heading(doc, "CONFRONTO TEMPORALE: PROGRAMMAZIONE VS REALTÀ")
        
        timeline_table = doc.add_table(rows=3, cols=2)
        timeline_table.style = 'Light Grid Accent 1'
        
        timeline_data = [
            ("DATA INIZIO PROGRAMMATO:", self._format_date(cantiere_info.get('data_inizio_prevista')),
             "DATA INIZIO EFFETTIVO:", self._format_date(cantiere_info.get('data_inizio_effettiva'))),
            ("DATA FINE PROGRAMMATO:", self._format_date(cantiere_info.get('data_fine_prevista')),
             "DATA FINE EFFETTIVO:", self._format_date(cantiere_info.get('data_fine_effettiva')))
        ]
        
        for row_idx, (label1, value1, label2, value2) in enumerate(timeline_data):
            row_cells = timeline_table.rows[row_idx].cells
            row_cells[0].text = f"{label1} {value1}"
            row_cells[1].text = f"{label2} {value2}"
            
            for cell_idx, cell in enumerate(row_cells):
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    if cell_idx == 0:
                        run.font.color.rgb = RGBColor(59, 130, 246)  # Blu per programmato
                    else:
                        run.font.color.rgb = RGBColor(34, 197, 94)  # Verde per effettivo
        
        # Riga per lo stato
        status_row = timeline_table.add_row().cells
        status_row[0].text = "STATO AVANZAMENTO:"
        status_row[1].text = f"{'CANTIERE ATTIVO' if cantiere_info.get('e_in_corso') else 'CANTIERE TERMINATO'}"
        
        for cell in status_row:
            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(10)
                run.font.bold = True

        doc.add_paragraph()

        # Sezione informazioni aggiuntive
        self._add_section_heading(doc, "INFORMAZIONI TECNICHE E GEOREFERENZIAZIONE")
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ("TIPOLOGIA INTERVENTO:", cantiere_info.get('tipologia_intervento', 'N/D')),
            ("AREA SPECIFICA:", cantiere_info.get('area_descrizione', 'N/D')),
            ("QUOTA ALTIMETRICA:", cantiere_info.get('quota', 'N/D')),
            ("COORDINATE GPS:", self._format_coordinates(cantiere_info.get('coordinate_lat'), cantiere_info.get('coordinate_lon')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row_cells = info_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            # Formattazione
            label_para = row_cells[0].paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = self.COLOR_ACCENT
            
            value_para = row_cells[1].paragraphs[0]
            if value_para.runs:
                value_run = value_para.runs[0]
                value_run.font.size = Pt(9)

        doc.add_paragraph()

    def _add_giornale_page(self, doc, giornale, num, total, cantiere_info):
        """Aggiunge una pagina completa di giornale con tutte le 11 sezioni"""
        
        # Header pagina
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(f"GIORNALE N. {num}/{total} - {self._format_date(giornale.get('data'))}")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = self.COLOR_HEADER_BG

        page_num = doc.add_paragraph()
        page_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_num_run = page_num.add_run(f"Pag. {num + 1}")
        page_num_run.font.size = Pt(8)
        page_num_run.font.italic = True
        page_num_run.font.color.rgb = self.COLOR_GREY

        doc.add_paragraph()

        # ===== SEZIONE 1: INFORMAZIONI GENERALI =====
        self._add_section_heading(doc, "1. INFORMAZIONI GENERALI")
        
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_rows = [
            ("Data:", self._format_date(giornale.get('data'))),
            ("Ora Inizio:", giornale.get('ora_inizio', 'N/D')),
            ("Ora Fine:", giornale.get('ora_fine', 'N/D')),
            ("Responsabile Scavo:", giornale.get('responsabile_scavo', 'N/D')),
            ("Compilatore:", giornale.get('compilatore', 'N/D')),
        ]

        for i, (label, value) in enumerate(info_rows):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            label_para = row.cells[0].paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(9)

        doc.add_paragraph()

        # ===== SEZIONE 2: CONDIZIONI METEOROLOGICHE =====
        self._add_section_heading(doc, "2. CONDIZIONI METEOROLOGICHE")
        
        meteo_p = doc.add_paragraph()
        if giornale.get('condizioni_meteo'):
            meteo_p.add_run("Condizioni: ").bold = True
            meteo_p.add_run(giornale['condizioni_meteo'].upper())
            meteo_p.add_run("\n")
        
        if giornale.get('temperatura'):
            temp_run = meteo_p.add_run("Temperatura: ")
            temp_run.bold = True
            temps = [f"Attuale: {giornale['temperatura']}°C"]
            if giornale.get('temperatura_min'):
                temps.append(f"Min: {giornale['temperatura_min']}°C")
            if giornale.get('temperatura_max'):
                temps.append(f"Max: {giornale['temperatura_max']}°C")
            meteo_p.add_run(", ".join(temps))
            meteo_p.add_run("\n")
        
        if giornale.get('note_meteo'):
            note_run = meteo_p.add_run("Note: ")
            note_run.bold = True
            meteo_p.add_run(giornale['note_meteo'])

        doc.add_paragraph()

        # ===== SEZIONE 3: DESCRIZIONE LAVORI =====
        self._add_section_heading(doc, "3. DESCRIZIONE LAVORI")
        
        if giornale.get('descrizione_lavori'):
            desc_p = doc.add_paragraph(giornale['descrizione_lavori'])
            for run in desc_p.runs:
                run.font.size = Pt(9)
        
        # Campi ICCD specialistici
        if giornale.get('area_intervento'):
            area_p = doc.add_paragraph()
            area_run = area_p.add_run("Area di intervento: ")
            area_run.bold = True
            area_p.add_run(giornale['area_intervento'])
            for run in area_p.runs:
                run.font.size = Pt(9)
        
        if giornale.get('saggio'):
            saggio_p = doc.add_paragraph()
            saggio_run = saggio_p.add_run("Saggio: ")
            saggio_run.bold = True
            saggio_p.add_run(giornale['saggio'])
            for run in saggio_p.runs:
                run.font.size = Pt(9)
        
        if giornale.get('obiettivi'):
            obj_p = doc.add_paragraph()
            obj_run = obj_p.add_run("Obiettivi: ")
            obj_run.bold = True
            obj_p.add_run(giornale['obiettivi'])
            for run in obj_p.runs:
                run.font.size = Pt(9)
        
        if giornale.get('interpretazione'):
            interp_p = doc.add_paragraph()
            interp_run = interp_p.add_run("Interpretazione: ")
            interp_run.bold = True
            interp_p.add_run(giornale['interpretazione'])
            for run in interp_p.runs:
                run.font.size = Pt(9)
        
        if giornale.get('modalita_lavorazioni'):
            mod_p = doc.add_paragraph()
            mod_run = mod_p.add_run("Modalità di lavorazione: ")
            mod_run.bold = True
            mod_p.add_run(giornale['modalita_lavorazioni'])
            for run in mod_p.runs:
                run.font.size = Pt(9)

        doc.add_paragraph()

        # ===== SEZIONE 4: RISORSE IMPIEGATE =====
        self._add_section_heading(doc, "4. RISORSE IMPIEGATE")
        
        # Operatori
        if giornale.get('operatori_presenti'):
            operatori_heading = doc.add_paragraph()
            op_run = operatori_heading.add_run("Operatori:")
            op_run.bold = True
            op_run.font.size = Pt(9)
            
            op_table = doc.add_table(rows=1, cols=4)
            op_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = op_table.rows[0].cells
            headers = ['Nome', 'Qualifica', 'Ore', 'Note']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_para = header_cells[i].paragraphs[0]
                header_run = header_para.runs[0]
                header_run.font.bold = True
                header_run.font.size = Pt(8)
            
            # Data rows
            for op in giornale['operatori_presenti']:
                row = op_table.add_row()
                row.cells[0].text = f"{op.get('nome', '')} {op.get('cognome', '')}"
                row.cells[1].text = op.get('qualifica', 'N/D')
                row.cells[2].text = str(op.get('ore_lavorate', '8'))
                row.cells[3].text = op.get('note_presenza', '')
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

        # Attrezzature
        if giornale.get('attrezzatura_utilizzata'):
            attr_p = doc.add_paragraph()
            attr_run = attr_p.add_run("Attrezzature: ")
            attr_run.bold = True
            attr_p.add_run(giornale['attrezzatura_utilizzata'])
            for run in attr_p.runs:
                run.font.size = Pt(9)
        
        # Mezzi
        if giornale.get('mezzi_utilizzati'):
            mezzi_p = doc.add_paragraph()
            mezzi_run = mezzi_p.add_run("Mezzi: ")
            mezzi_run.bold = True
            mezzi_p.add_run(giornale['mezzi_utilizzati'])
            for run in mezzi_p.runs:
                run.font.size = Pt(9)

        doc.add_paragraph()

        # ===== SEZIONE 5: UNITÀ STRATIGRAFICHE =====
        us_list = giornale.get('us_elaborate', []) or []
        usm_list = giornale.get('usm_elaborate', []) or []
        usr_list = giornale.get('usr_elaborate', []) or []
        
        if us_list or usm_list or usr_list:
            self._add_section_heading(doc, "5. UNITÀ STRATIGRAFICHE ELABORATE")
            
            if us_list:
                p = doc.add_paragraph()
                p_run = p.add_run("US: ")
                p_run.bold = True
                p.add_run(', '.join(str(u) for u in us_list))
            
            if usm_list:
                p = doc.add_paragraph()
                p_run = p.add_run("USM: ")
                p_run.bold = True
                p.add_run(', '.join(str(u) for u in usm_list))
            
            if usr_list:
                p = doc.add_paragraph()
                p_run = p.add_run("USR: ")
                p_run.bold = True
                p.add_run(', '.join(str(u) for u in usr_list))

            doc.add_paragraph()

        # ===== SEZIONE 6: MATERIALI RINVENUTI =====
        if giornale.get('materiali_rinvenuti'):
            self._add_section_heading(doc, "6. MATERIALI RINVENUTI")
            p = doc.add_paragraph(giornale['materiali_rinvenuti'])
            for run in p.runs:
                run.font.size = Pt(9)
            doc.add_paragraph()

        # ===== SEZIONE 7: DOCUMENTAZIONE PRODOTTA =====
        if giornale.get('documentazione_prodotta'):
            self._add_section_heading(doc, "7. DOCUMENTAZIONE PRODOTTA")
            p = doc.add_paragraph(giornale['documentazione_prodotta'])
            for run in p.runs:
                run.font.size = Pt(9)
            doc.add_paragraph()

        # ===== SEZIONE 8: DISPOSIZIONI E ORDINI =====
        disposizioni = []
        if giornale.get('disposizioni_rup'):
            disposizioni.append(("RUP", giornale['disposizioni_rup']))
        if giornale.get('disposizioni_direttore'):
            disposizioni.append(("Direttore Lavori", giornale['disposizioni_direttore']))
        
        if disposizioni:
            self._add_section_heading(doc, "8. DISPOSIZIONI E ORDINI")
            for label, val in disposizioni:
                p = doc.add_paragraph()
                p_run = p.add_run(f"{label}: ")
                p_run.bold = True
                p.add_run(val)
                for run in p.runs:
                    run.font.size = Pt(9)
            doc.add_paragraph()

        # ===== SEZIONE 9: EVENTI PARTICOLARI =====
        eventi = []
        if giornale.get('sospensioni'):
            eventi.append(("Sospensioni", giornale['sospensioni']))
        if giornale.get('contestazioni'):
            eventi.append(("Contestazioni", giornale['contestazioni']))
        if giornale.get('incidenti'):
            eventi.append(("Incidenti", giornale['incidenti']))
        if giornale.get('problematiche'):
            eventi.append(("Problematiche", giornale['problematiche']))
        
        if eventi:
            self._add_section_heading(doc, "9. EVENTI PARTICOLARI")
            for label, val in eventi:
                p = doc.add_paragraph()
                p_run = p.add_run(f"{label}: ")
                p_run.bold = True
                p.add_run(val)
                for run in p.runs:
                    run.font.size = Pt(9)
            doc.add_paragraph()

        # ===== SEZIONE 10: NOTE E OSSERVAZIONI =====
        if giornale.get('note_generali') or giornale.get('sopralluoghi'):
            self._add_section_heading(doc, "10. NOTE E OSSERVAZIONI")
            
            if giornale.get('note_generali'):
                p = doc.add_paragraph(giornale['note_generali'])
                for run in p.runs:
                    run.font.size = Pt(9)
            
            if giornale.get('sopralluoghi'):
                p = doc.add_paragraph()
                p_run = p.add_run("Sopralluoghi: ")
                p_run.bold = True
                p.add_run(giornale['sopralluoghi'])
                for run in p.runs:
                    run.font.size = Pt(9)
            
            doc.add_paragraph()

        # ===== SEZIONE 11: VALIDAZIONE =====
        self._add_section_heading(doc, "11. STATO VALIDAZIONE")
        
        val_table = doc.add_table(rows=4, cols=2)
        val_table.style = 'Light Grid Accent 1'

        val_rows = [
            ("Validato:", "✓ SI" if giornale.get('validato') else "✗ NO"),
            ("Data Validazione:", giornale.get('data_validazione', 'N/D')),
            ("Data Creazione:", giornale.get('created_at', 'N/D')),
            ("Ultimo Aggiornamento:", giornale.get('updated_at', 'N/D')),
        ]

        for i, (label, value) in enumerate(val_rows):
            row = val_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            label_para = row.cells[0].paragraphs[0]
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(8)

        doc.add_paragraph()

        # ===== SEZIONE 12: DOCUMENTAZIONE FOTOGRAFICA =====
        foto_list = giornale.get('foto', [])
        if foto_list:
            self._add_section_heading(doc, "12. DOCUMENTAZIONE FOTOGRAFICA")
            
            # Conteggio foto
            p = doc.add_paragraph()
            p_run = p.add_run(f"Numero foto collegate: {len(foto_list)}")
            p_run.bold = True
            p_run.font.size = Pt(9)
            
            doc.add_paragraph()
            
            # Aggiungi ogni foto
            for idx, foto in enumerate(foto_list, 1):
                try:
                    # Titolo/didascalia foto
                    title = foto.get('title') or foto.get('filename') or f'Foto {idx}'
                    foto_heading = doc.add_paragraph()
                    foto_heading_run = foto_heading.add_run(f"Foto {idx}: {title}")
                    foto_heading_run.bold = True
                    foto_heading_run.font.size = Pt(10)
                    foto_heading_run.font.color.rgb = self.COLOR_ACCENT
                    
                    # Inserisci immagine se disponibile
                    image_bytes = foto.get('_image_bytes')
                    if image_bytes:
                        try:
                            image_stream = io.BytesIO(image_bytes)
                            doc.add_picture(image_stream, width=Inches(5.0))
                            image_stream.close()
                        except Exception as e:
                            logger.warning(f"Could not insert image {idx}: {e}")
                            doc.add_paragraph(f"[Errore inserimento immagine: {str(e)}]")
                    else:
                        doc.add_paragraph("[Immagine non disponibile]")
                    
                    # Descrizione aggiuntiva se presente
                    if foto.get('description'):
                        desc_p = doc.add_paragraph()
                        desc_run = desc_p.add_run(foto['description'])
                        desc_run.font.italic = True
                        desc_run.font.size = Pt(8)
                        desc_run.font.color.rgb = self.COLOR_GREY
                    
                    doc.add_paragraph()  # Spacer
                    
                except Exception as e:
                    logger.error(f"Error adding photo {idx} to Word: {e}")
                    doc.add_paragraph(f"[Errore inserimento foto {idx}: {str(e)}]")

    def _add_signature_page(self, doc, cantiere_info, site_info):
        """Pagina finale con firme"""
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("FIRME E VALIDAZIONI")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = self.COLOR_HEADER_BG

        doc.add_paragraph()

        firme_text = (
            "Sottoscritti il presente Giornale di Cantiere:\n\n"
            "Il Responsabile di Scavo: ____________________________     Data: __________\n"
            "Nome: _________________________________ Qualifica: _______________________\n\n\n"
            "Il Direttore dei Lavori: ____________________________     Data: __________\n"
            "Nome: _________________________________ Qualifica: _______________________\n\n\n"
            "Il Responsabile del Procedimento: ____________________________     Data: __________\n"
            "Nome: _________________________________ Qualifica: _______________________\n\n\n"
            "Il Rappresentante della Committenza: ____________________________     Data: __________\n"
            "Nome: _________________________________ Qualifica: _______________________"
        )
        
        p = doc.add_paragraph(firme_text)
        for run in p.runs:
            run.font.size = Pt(9)

        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            f"Documento generato da FastZoom Archaeological System\n"
            f"Data: {datetime.now().strftime('%d/%m/%Y ore %H:%M:%S')}\n"
            f"Sito: {site_info.get('name', 'N/D')}\n"
            f"Cantiere: {cantiere_info.get('nome_completo', cantiere_info.get('nome', 'N/D'))}\n"
            f"Stato: {cantiere_info.get('stato_formattato', 'N/D')} | "
            f"Durata: {cantiere_info.get('durata_giorni', 'N/D')} giorni | "
            f"Priorità: {cantiere_info.get('priorita', 'N/D')}/5"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = self.COLOR_GREY

    def _add_section_heading(self, doc, text):
        """Aggiunge heading di sezione"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(17.8)
        cell = table.rows[0].cells[0]
        cell.text = ""
        self._set_cell_shading(cell, "E8EEF5")
        self._set_cell_borders(cell, "2C5AA0", size="8")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_ACCENT

    def _format_date(self, date_value) -> str:
        """Formatta data italiana"""
        if not date_value:
            return 'N/D'
        
        if isinstance(date_value, str):
            try:
                dt = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
                return dt.strftime('%d/%m/%Y')
            except:
                return date_value
        
        try:
            return date_value.strftime('%d/%m/%Y')
        except:
            return str(date_value)

    def _format_coordinates(self, lat, lon) -> str:
        """Formatta coordinate GPS in formato leggibile"""
        if not lat or not lon:
            return 'N/D'
        
        try:
            # Formatta coordinate con precisione decimale
            lat_formatted = f"{float(lat):.6f}"
            lon_formatted = f"{float(lon):.6f}"
            return f"{lat_formatted}°N, {lon_formatted}°E"
        except (ValueError, TypeError):
            return f"{lat}, {lon}" if lat and lon else 'N/D'


# Istanza globale
if DOCX_AVAILABLE:
    _word_generator = GiornaleWordGeneratorV2()

    def generate_giornale_word_quick(giornali: List[Dict[str, Any]],
                                    cantiere_info: Dict[str, Any],
                                    site_info: Dict[str, Any]) -> bytes:
        """Funzione di utilità - Genera Word rapidamente"""
        return _word_generator.generate_giornale_word(giornali, cantiere_info, site_info)
else:
    def generate_giornale_word_quick(*args, **kwargs) -> bytes:
        """Stub quando docx non disponibile"""
        raise ImportError("python-docx required: pip install python-docx")
