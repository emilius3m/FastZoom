# app/services/giornale_word_export.py
"""
Export Giornale di Cantiere compilando direttamente template Word esistente
Mantiene layout IDENTICO al documento MiC originale
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date
from typing import Optional, Dict, Any, List
import os
from pathlib import Path
import re
import io
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class GiornaleWordExporter:
    """
    Compila template Word esistente con dati Giornale di Cantiere
    NON crea nuovo template - modifica l'originale mantenendo il layout
    """

    def __init__(self, template_path: str):
        """
        Args:
            template_path: Path al file .docx template originale
        """
        self.template_path = template_path

    def export_giornali_list(self, export_data: Dict[str, Any], output_path: str) -> str:
        """
        Compila report lista giornali (Registro Completo)
        """
        doc = Document(self.template_path)
        self._clear_body(doc)

        # A. INTESTAZIONE E RIEPILOGO
        self._add_site_header(doc, export_data.get('site_info', {}))
        
        doc.add_heading('REGISTRO GIORNALE DI CANTIERE', 0)
        
        # Metadati Export
        self._add_section_title(doc, 'METADATI EXPORT')
        export_metadata = export_data.get('export_metadata', {})
        self._add_kv(doc, "Data Export", self._format_date(export_metadata.get('export_date')))
        self._add_kv(doc, "Utente Export", export_metadata.get('user', ''))
        if export_metadata.get('filters'):
            self._add_kv(doc, "Filtri Applicati", export_metadata.get('filters', ''))

        # Statistiche
        stats = export_data.get('stats', {})
        self._add_section_title(doc, 'STATISTICHE RIEPILOGATIVE')
        self._add_kv(doc, "Totale Giornali", str(stats.get('total_giornali', 0)))
        self._add_kv(doc, "Giornali Validati", str(stats.get('validated_giornali', 0)))
        self._add_kv(doc, "Giornali Pendenti", str(stats.get('pending_giornali', 0)))
        self._add_kv(doc, "Operatori Attivi", str(stats.get('operatori_attivi', 0)))
        self._add_kv(doc, "Percentuale Completamento", f"{stats.get('validation_percentage', 0)}%")
        
        doc.add_paragraph() # Spacer

        # B. TABELLA RIEPILOGATIVA
        self._add_section_title(doc, 'ELENCO GIORNALI')
        self._create_grouped_giornali_table(doc, export_data.get('giornali', []))

        # C. DETTAGLIO GIORNALI (Schede singole)
        giornali = export_data.get('giornali', [])
        # Ordina cronologicamente
        giornali.sort(key=lambda x: x.get('data') or '')

        for giornale in giornali:
            doc.add_page_break()
            self._add_giornale_detail(doc, giornale, is_list=True)

        doc.save(output_path)
        return output_path

    def export_single_giornale(self, giornale_data: Dict[str, Any], output_path: str) -> str:
        """
        Compila scheda singolo giornale
        """
        doc = Document(self.template_path)
        self._clear_body(doc)

        # Intestazione Sito (Opzionale, se vogliamo mantenere coerenza con list)
        self._add_site_header(doc, giornale_data.get('site_info', {}))

        # Dettaglio Giornale
        self._add_giornale_detail(doc, giornale_data)

        doc.save(output_path)
        return output_path

    # ===== HELPER METODI PER LAYOUT =====

    def _add_site_header(self, doc: Document, site_info: Dict[str, Any]):
        """Aggiunge intestazione sito"""
        self._add_section_title(doc, 'INFORMAZIONI SITO')
        self._add_kv(doc, "Nome Sito", site_info.get('name', ''))
        self._add_kv(doc, "Codice Sito", site_info.get('code', ''))
        self._add_kv(doc, "Località", site_info.get('location', ''))
        doc.add_paragraph()

    def _add_giornale_detail(self, doc: Document, g: Dict[str, Any], is_list: bool = False):
        """Genera la scheda dettagliata di un giornale"""
        
        # Titolo
        title = f"DETTAGLIO GIORNALE - {self._format_date(g.get('data'))}"
        doc.add_heading(title, level=1)

        # Info Base (Layout: Label bold, poi valore)
        p = doc.add_paragraph()
        run = p.add_run("DATA GIORNALE: ")
        run.bold = True
        p.add_run(self._format_date(g.get('data')))
        
        # Orari inline? O riga per riga? "ORA INIZIO: ..." User req: riga per riga
        self._add_kv(doc, "ORA INIZIO", self._format_time(g.get('ora_inizio')))
        self._add_kv(doc, "ORA FINE", self._format_time(g.get('ora_fine')))
        self._add_kv(doc, "COMPILATORE", g.get('compilatore', ''))
        self._add_kv(doc, "RESPONSABILE SCAVO", g.get('responsabile_scavo', ''))
        doc.add_paragraph()

        # Condizioni Operative
        self._add_section_title(doc, "CONDIZIONI OPERATIVE")
        self._add_kv(doc, "Condizioni Meteo", g.get('condizioni_meteo', ''))
        if g.get('temperatura_min'):
            self._add_kv(doc, "Temperatura Minima", f"{g.get('temperatura_min')}°C")
        if g.get('temperatura_max'):
            self._add_kv(doc, "Temperatura Massima", f"{g.get('temperatura_max')}°C")
        if g.get('note_meteo'):
            self._add_kv(doc, "Note Meteo", g.get('note_meteo', ''))
        
        # Descrizione Lavori
        self._add_section_title(doc, "DESCRIZIONE LAVORI")
        self._add_kv(doc, "Descrizione Lavori", g.get('descrizione_lavori', ''))
        
        # Nuovi campi ICCD
        if g.get('area_intervento'): self._add_kv(doc, "Area Intervento", g.get('area_intervento', ''))
        if g.get('saggio'): self._add_kv(doc, "Saggio", g.get('saggio', ''))
        if g.get('obiettivi'): self._add_kv(doc, "Obiettivi", g.get('obiettivi', ''))
        if g.get('interpretazione'): self._add_kv(doc, "Interpretazione", g.get('interpretazione', ''))

        self._add_kv(doc, "Modalità Lavorazioni", g.get('modalita_lavorazioni', ''))
        self._add_kv(doc, "Attrezzatura Utilizzata", g.get('attrezzatura_utilizzata', ''))
        self._add_kv(doc, "Mezzi Utilizzati", g.get('mezzi_utilizzati', ''))

        # Documentazione Archeologica
        self._add_section_title(doc, "DOCUMENTAZIONE ARCHEOLOGICA")
        
        us = g.get('us_elaborate', [])
        if us: self._add_kv(doc, "US Elaborate", us if isinstance(us, str) else ", ".join(us))
        
        usm = g.get('usm_elaborate', [])
        if usm: self._add_kv(doc, "USM Elaborate", usm if isinstance(usm, str) else ", ".join(usm))
        
        usr = g.get('usr_elaborate', [])
        if usr: self._add_kv(doc, "USR Elaborate", usr if isinstance(usr, str) else ", ".join(usr))

        self._add_kv(doc, "Materiali Rinvenuti", g.get('materiali_rinvenuti', ''))
        self._add_kv(doc, "Campioni Prelevati", g.get('campioni_prelevati', ''))
        self._add_kv(doc, "Documentazione Prodotta", g.get('documentazione_prodotta', ''))
        self._add_kv(doc, "Strutture", g.get('strutture', ''))

        # Operatori - Tabella
        self._add_section_title(doc, "OPERATORI PRESENTI")
        operatori = g.get('operatori_presenti', [])
        if operatori:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Nome"
            hdr[1].text = "Qualifica"
            hdr[2].text = "Ruolo"
            for c in hdr: c.paragraphs[0].runs[0].bold = True
            
            for op in operatori:
                row = table.add_row()
                row.cells[0].text = f"{op.get('nome', '')} {op.get('cognome', '')}"
                row.cells[1].text = op.get('qualifica', '')
                row.cells[2].text = op.get('ruolo', '')
        else:
            doc.add_paragraph("Nessun operatore registrato.")
        
        doc.add_paragraph()

        # Sopralluoghi e Disposizioni
        self._add_section_title(doc, "SOPRALLUOGHI E DISPOSIZIONI")
        self._add_kv(doc, "Sopralluoghi", g.get('sopralluoghi', ''))
        self._add_kv(doc, "Disposizioni RUP", g.get('disposizioni_rup', ''))
        self._add_kv(doc, "Disposizioni Direttore", g.get('disposizioni_direttore', ''))

        # Eventi Particolari
        self._add_section_title(doc, "EVENTI PARTICOLARI")
        self._add_kv(doc, "Contestazioni", g.get('contestazioni', ''))
        self._add_kv(doc, "Sospensioni", g.get('sospensioni', ''))
        self._add_kv(doc, "Incidenti", g.get('incidenti', ''))
        self._add_kv(doc, "Forniture", g.get('forniture', ''))

        # Note
        self._add_section_title(doc, "NOTE E PROBLEMATICHE")
        self._add_kv(doc, "Note Generali", g.get('note_generali', ''))
        self._add_kv(doc, "Problematiche", g.get('problematiche', ''))

        # Stato Validazione
        self._add_section_title(doc, "STATO VALIDAZIONE")
        status_text = 'Validato' if g.get('validato') else 'In Attesa'
        self._add_kv(doc, "Stato Validazione", status_text)
        if g.get('data_validazione'):
            self._add_kv(doc, "Data Validazione", self._format_datetime(g.get('data_validazione')))

        # Firme
        doc.add_paragraph()
        self._add_section_title(doc, "FIRME")
        table_firme = doc.add_table(rows=1, cols=3) # Layout invisibile per firme
        # width adjustments needed? for now standard
        c = table_firme.rows[0].cells
        
        def add_sign_line(cell, label):
            p = cell.add_paragraph()
            p.add_run("_" * 20)
            cell.add_paragraph(label)
        
        add_sign_line(c[0], "Responsabile Scavo")
        add_sign_line(c[1], "Direttore Lavori")
        add_sign_line(c[2], "RUP")

        # Foto
        self._add_photos_section(doc, g.get('foto', []))

    def _add_section_title(self, doc: Document, text: str):
        """Heading 2 uppercase"""
        doc.add_heading(text.upper(), level=2)

    def _add_kv(self, doc: Document, key: str, value: str):
        """Aggiunge paragrafo 'Key: Value'"""
        # Se value è vuoto, mostriamo comunque la label? L'utente ha messo label vuote nell'esempio. Sì.
        p = doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        run_key.bold = True
        p.add_run(str(value) if value else "")

    def _clear_body(self, doc: Document):
        """Svuota il corpo del documento mantenendo header/footer e section properties"""
        body = doc.element.body
        # Rimuovi tutti gli elementi tranne sectPr (che contiene settings sezione e header/footer refs)
        for element in list(body):
            if element.tag.endswith('sectPr'):
                continue
            body.remove(element)

    def _create_grouped_giornali_table(self, doc: Document, giornali: List[Dict[str, Any]]):
        """Crea tabella giornali raggruppata per mese"""
        if not giornali:
            doc.add_paragraph("Nessun giornale trovato.")
            return

        # Ordina per data (decrescente o crescente)
        # Register standard: Cronologico (Crescente)
        giornali.sort(key=lambda x: x.get('data') or '', reverse=False)

        # Raggruppa per mese
        from itertools import groupby
        def get_month_year(g):
            d = g.get('data')
            if isinstance(d, str):
                try:
                    return datetime.fromisoformat(d.replace('Z', '')).strftime('%B %Y')
                except: return "Data sconosciuta"
            return d.strftime('%B %Y') if d else "Data sconosciuta"

        table = doc.add_table(rows=0, cols=6)
        table.style = 'Table Grid'
        
        # Header Colonne
        headers = ['Data', 'Orari', 'Responsabile', 'Meteo', 'Stato', 'Note']
        
        for key, group in groupby(giornali, key=get_month_year):
            # Header Gruppo (Mese)
            row_group = table.add_row()
            # Merge calle per header gruppo
            row_group.cells[0].merge(row_group.cells[-1])
            row_group.cells[0].text = key.upper()
            row_group.cells[0].paragraphs[0].runs[0].bold = True
            shading_elm = parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w')))
            row_group.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

            # Header Colonne
            row_header = table.add_row()
            for i, h in enumerate(headers):
                row_header.cells[i].text = h
                row_header.cells[i].paragraphs[0].runs[0].bold = True

            # Dati
            for g in group:
                row = table.add_row()
                row.cells[0].text = self._format_date(g.get('data'))
                
                ora_start = self._format_time(g.get('ora_inizio'))
                ora_end = self._format_time(g.get('ora_fine'))
                row.cells[1].text = f"{ora_start}\n{ora_end}"
                
                row.cells[2].text = g.get('responsabile_scavo', '') or g.get('compilatore', '')
                row.cells[3].text = g.get('condizioni_meteo', '')
                row.cells[4].text = 'Validato' if g.get('validato') else 'In Attesa'
                
                note = g.get('descrizione_lavori', '') or g.get('note_generali', '')
                if len(note) > 100: note = note[:100] + "..."
                row.cells[5].text = note

    def _add_photos_section(self, doc: Document, photos: List[Dict[str, Any]]):
        """
        Aggiunge sezione fotografica alla fine del documento
        """
        if not photos:
            return

        doc.add_page_break()
        doc.add_heading('DOCUMENTAZIONE FOTOGRAFICA', level=1)
        
        # Aggiungi conteggio
        p = doc.add_paragraph()
        p.add_run(f"Foto collegate: {len(photos)}").bold = True
        
        # Aggiungi foto
        for photo in photos:
            try:
                # Titolo foto
                title = photo.get('title') or photo.get('original_filename') or 'Foto'
                doc.add_heading(title, level=3)
                
                # Image bytes must be pre-loaded in '_image_bytes' key
                image_bytes = photo.get('_image_bytes')
                
                if image_bytes:
                    # Save bytes to temp file because add_picture needs path or stream
                    image_stream = io.BytesIO(image_bytes)
                    doc.add_picture(image_stream, width=Inches(6.0))
                else:
                    doc.add_paragraph("[Immagine non disponibile]")
                
                # Didascalia/Descrizione
                if photo.get('description'):
                    p = doc.add_paragraph(photo.get('description'))
                    p.italic = True
                
                doc.add_paragraph("") # Spacer
                
            except Exception as e:
                doc.add_paragraph(f"[Errore inserimento foto: {str(e)}]")

    # ===== METODI HELPER PER MANIPOLAZIONE DOCUMENTO =====

    def _replace_text(self, doc: Document, placeholder: str, value: str):
        """
        Sostituisce placeholder in TUTTO il documento (paragrafi, tabelle, header, footer)
        Mantiene formattazione originale
        """
        value = str(value) if value is not None else ''

        # Sostituisci in paragrafi
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                self._replace_in_paragraph(paragraph, placeholder, value)

        # Sostituisci in tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, value)

        # Sostituisci in header/footer
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                if placeholder in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder, value)

            # Footer
            for paragraph in section.footer.paragraphs:
                if placeholder in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder, value)

    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str):
        """
        Sostituisce testo nel paragrafo mantenendo TUTTA la formattazione originale
        """
        # IMPORTANTE: Sostituisci nel testo completo mantenendo i run
        if placeholder not in paragraph.text:
            return

        # Trova tutti i run che contengono parte del placeholder
        full_text = ''
        run_texts = []

        for run in paragraph.runs:
            run_texts.append(run.text)
            full_text += run.text

        # Se il placeholder è presente, sostituisci
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, value)

            # Cancella tutti i run esistenti tranne il primo
            for _ in range(len(paragraph.runs) - 1):
                paragraph._element.remove(paragraph.runs[-1]._element)

            # Aggiorna il primo run con il testo nuovo
            if paragraph.runs:
                paragraph.runs[0].text = new_text

    def _compile_giornali_table(self, doc: Document, giornali: List[Dict[str, Any]]):
        """
        Compila tabella con elenco giornali
        """
        # Trova tabella giornali nel documento (cerca placeholder specifico)
        for table in doc.tables:
            # Cerca prima cella con "GIORNALE" o simile
            if self._is_giornali_table(table):
                self._fill_giornali_table(table, giornali)
                break

    def _is_giornali_table(self, table) -> bool:
        """Identifica se è la tabella giornali"""
        # Controlla se prima riga contiene intestazioni giornali
        if len(table.rows) > 0:
            first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
            return 'DATA' in first_row_text.upper() or 'GIORNALE' in first_row_text.upper()
        return False

    def _fill_giornali_table(self, table, giornali: List[Dict[str, Any]]):
        """
        Compila celle tabella giornali con dati
        """
        # Se non ci sono giornali, aggiungi riga vuota
        if not giornali:
            if len(table.rows) > 1:
                # Mantieni solo header
                for _ in range(len(table.rows) - 1):
                    table._tbl.remove(table.rows[-1]._tr)
            return

        # Rimuovi righe esistenti tranne header
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # Aggiungi riga per ogni giornale
        for giornale in giornali:
            row = table.add_row()
            
            # Data
            row.cells[0].text = self._format_date(giornale.get('data'))
            
            # Orari
            ora_inizio = self._format_time(giornale.get('ora_inizio'))
            ora_fine = self._format_time(giornale.get('ora_fine'))
            row.cells[1].text = f"{ora_inizio} - {ora_fine}" if ora_inizio and ora_fine else ""
            
            # Responsabile
            row.cells[2].text = giornale.get('responsabile_scavo', '')
            
            # Condizioni meteo
            row.cells[3].text = giornale.get('condizioni_meteo', '')
            
            # Stato
            row.cells[4].text = 'Validato' if giornale.get('validato') else 'In Attesa'
            
            # Note (troncate)
            note = giornale.get('note_generali', '')
            if note and len(note) > 50:
                note = note[:50] + "..."
            row.cells[5].text = note

    def _compile_operatori_table(self, doc: Document, operatori: List[Dict[str, Any]]):
        """
        Compila tabella con elenco operatori presenti
        """
        # Trova tabella operatori nel documento
        for table in doc.tables:
            # Cerca prima cella con "OPERATORI" o simile
            if self._is_operatori_table(table):
                self._fill_operatori_table(table, operatori)
                break

    def _is_operatori_table(self, table) -> bool:
        """Identifica se è la tabella operatori"""
        if len(table.rows) > 0:
            first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
            return 'OPERATORI' in first_row_text.upper() or 'NOME' in first_row_text.upper()
        return False

    def _fill_operatori_table(self, table, operatori: List[Dict[str, Any]]):
        """
        Compila celle tabella operatori con dati
        """
        # Se non ci sono operatori, aggiungi riga vuota
        if not operatori:
            if len(table.rows) > 1:
                # Mantieni solo header
                for _ in range(len(table.rows) - 1):
                    table._tbl.remove(table.rows[-1]._tr)
            return

        # Rimuovi righe esistenti tranne header
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # Aggiungi riga per ogni operatore
        for operatore in operatori:
            row = table.add_row()
            
            # Nome completo
            row.cells[0].text = f"{operatore.get('nome', '')} {operatore.get('cognome', '')}"
            
            # Qualifica
            row.cells[1].text = operatore.get('qualifica', '')
            
            # Ruolo
            row.cells[2].text = operatore.get('ruolo', '')

    def _format_date(self, date_value) -> str:
        """Formatta data in formato italiano"""
        if not date_value:
            return ''

        if isinstance(date_value, str):
            try:
                date_obj = datetime.fromisoformat(date_value.replace('Z', '+00:00')).date()
                return date_obj.strftime('%d/%m/%Y')
            except:
                return date_value

        try:
            return date_value.strftime('%d/%m/%Y')
        except:
            return str(date_value)

    def _format_time(self, time_value) -> str:
        """Formatta ora"""
        if not time_value:
            return ''

        if isinstance(time_value, str):
            return time_value

        try:
            return time_value.strftime('%H:%M')
        except:
            return str(time_value)

    def _format_datetime(self, datetime_value) -> str:
        """Formatta data e ora"""
        if not datetime_value:
            return ''

        if isinstance(datetime_value, str):
            try:
                dt_obj = datetime.fromisoformat(datetime_value.replace('Z', '+00:00'))
                return dt_obj.strftime('%d/%m/%Y %H:%M')
            except:
                return datetime_value

        try:
            return datetime_value.strftime('%d/%m/%Y %H:%M')
        except:
            return str(datetime_value)


# ===== FUNZIONE HELPER PER FASTAPI =====

def create_giornale_word_from_template(export_data: Dict[str, Any], template_docx_path: str, output_dir: str) -> str:
    """
    Crea documento Word compilato da dati giornale

    Args:
        export_data: Dizionario con dati di export
        template_docx_path: Path al template .docx originale
        output_dir: Directory dove salvare output

    Returns:
        Path del file generato
    """
    exporter = GiornaleWordExporter(template_docx_path)

    # Genera nome file output
    site_info = export_data.get('site_info', {})
    site_name = site_info.get('name', 'Sito').replace(' ', '_').replace(',', '')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if export_data.get('single_giornale'):
        # Export singolo giornale
        giornale_data = export_data.get('giornale', {})
        data_giornale = exporter._format_date(giornale_data.get('data')).replace('/', '-')
        filename = f"Giornale_{site_name}_{data_giornale}_{timestamp}.docx"
    else:
        # Export lista giornali
        filename = f"Giornali_{site_name}_{timestamp}.docx"
    
    output_path = os.path.join(output_dir, filename)

    # Compila template
    if export_data.get('single_giornale'):
        return exporter.export_single_giornale(export_data.get('giornale', {}), output_path)
    else:
        return exporter.export_giornali_list(export_data, output_path)