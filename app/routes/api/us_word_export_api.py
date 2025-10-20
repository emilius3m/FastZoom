# app/routes/api/us_word_export_api.py - VERSIONE CORRETTA CON TEMPLATE

"""
API per esportazione schede US/USM in formato Word
Usa template .docx con placeholder compilati direttamente
"""

from typing import List, Dict, Any, Optional
from uuid import UUID
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from loguru import logger
import io
import zipfile
from datetime import datetime
from docx import Document

from app.database.db import get_async_session
from app.core.security import get_current_user_id_with_blacklist, get_current_user_sites_with_blacklist
from app.models.stratigraphy import UnitaStratigrafica, UnitaStratigraficaMuraria

router = APIRouter(prefix="/api/us-export", tags=["us-export"])

# Path ai template .docx con placeholder
TEMPLATES_DIR = Path("app/templates/word")
US_TEMPLATE_PATH = TEMPLATES_DIR / "US_Template_con_Placeholder.docx"
USM_TEMPLATE_PATH = TEMPLATES_DIR / "USM_Template_con_Placeholder.docx"  # Da creare


# ===== FUNZIONI HELPER =====

async def verify_site_access(site_id: UUID, user_sites: List[Dict[str, Any]]) -> bool:
    """Verifica accesso utente al sito"""
    return any(s["id"] == str(site_id) for s in user_sites)


def compile_us_template(us: UnitaStratigrafica, template_path: Path) -> bytes:
    """
    Compila template Word con dati US
    Sostituisce placeholder {{...}} con dati database
    """
    # Carica template
    doc = Document(template_path)

    # Prepara dati per sostituzione
    replacements = {
        '{{us_code}}': us.us_code or '',
        '{{ente_responsabile}}': us.ente_responsabile or '',
        '{{anno}}': str(us.anno) if us.anno else '',
        '{{ufficio_mic}}': us.ufficio_mic or '',
        '{{identificativo_rif}}': us.identificativo_rif or '',
        '{{localita}}': us.localita or '',
        '{{area_struttura}}': us.area_struttura or '',
        '{{ambiente_unita}}': us.ambiente_unita_funzione or '',
        '{{saggio}}': us.saggio or '',
        '{{posizione}}': us.posizione or '',
        '{{settori}}': us.settori or '',
        '{{piante}}': us.piante_riferimenti or '',
        '{{prospetti}}': us.prospetti_riferimenti or '',
        '{{sezioni}}': us.sezioni_riferimenti or '',
        '{{fotografie}}': '',  # Non esiste nel modello
        '{{rif_tabelle}}': '',  # Non esiste nel modello
        '{{definizione}}': us.definizione or '',
        '{{criteri_distinzione}}': us.criteri_distinzione or '',
        '{{modo_formazione}}': us.modo_formazione or '',
        '{{inorganici}}': us.componenti_inorganici or '',
        '{{organici}}': us.componenti_organici or '',
        '{{consistenza}}': us.consistenza or '',
        '{{colore}}': us.colore or '',
        '{{misure}}': us.misure or '',
        '{{conservazione}}': us.stato_conservazione or '',

        # Matrix Harris - converti JSON in testo
        '{{uguale_a}}': ', '.join(map(str, us.sequenza_fisica.get('uguale_a', []))) if us.sequenza_fisica else '',
        '{{si_lega_a}}': ', '.join(map(str, us.sequenza_fisica.get('si_lega_a', []))) if us.sequenza_fisica else '',
        '{{gli_si_appoggia}}': ', '.join(
            map(str, us.sequenza_fisica.get('gli_si_appoggia', []))) if us.sequenza_fisica else '',
        '{{si_appoggia_a}}': ', '.join(
            map(str, us.sequenza_fisica.get('si_appoggia_a', []))) if us.sequenza_fisica else '',
        '{{coperto_da}}': ', '.join(map(str, us.sequenza_fisica.get('coperto_da', []))) if us.sequenza_fisica else '',
        '{{copre}}': ', '.join(map(str, us.sequenza_fisica.get('copre', []))) if us.sequenza_fisica else '',
        '{{tagliato_da}}': ', '.join(map(str, us.sequenza_fisica.get('tagliato_da', []))) if us.sequenza_fisica else '',
        '{{taglia}}': ', '.join(map(str, us.sequenza_fisica.get('taglia', []))) if us.sequenza_fisica else '',
        '{{riempito_da}}': ', '.join(map(str, us.sequenza_fisica.get('riempito_da', []))) if us.sequenza_fisica else '',
        '{{riempie}}': ', '.join(map(str, us.sequenza_fisica.get('riempie', []))) if us.sequenza_fisica else '',
        '{{posteriore_a}}': ', '.join(
            map(str, us.sequenza_fisica.get('posteriore_a', []))) if us.sequenza_fisica else '',
        '{{anteriore_a}}': ', '.join(map(str, us.sequenza_fisica.get('anteriore_a', []))) if us.sequenza_fisica else '',

        '{{descrizione}}': us.descrizione or '',
        '{{osservazioni}}': us.osservazioni or '',
        '{{interpretazione}}': us.interpretazione or '',
        '{{datazione}}': us.datazione or '',
        '{{periodo}}': us.periodo or '',
        '{{fase}}': us.fase or '',
        '{{attivita}}': '',  # Non esiste nel modello
        '{{elementi_datanti}}': us.elementi_datanti or '',
        '{{reperti}}': us.dati_quantitativi_reperti or '',

        # Checkbox - simboli
        '{{flottazione}}': '☑' if (us.campionature and us.campionature.get('flottazione')) else '☐',
        '{{setacciatura}}': '☑' if (us.campionature and us.campionature.get('setacciatura')) else '☐',

        '{{affidabilita}}': us.affidabilita_stratigrafica or '',
        '{{resp_scientifico}}': us.responsabile_scientifico or '',
        '{{data_rilevamento}}': us.data_rilevamento.strftime('%d/%m/%Y') if us.data_rilevamento else '',
        '{{resp_compilazione}}': us.responsabile_compilazione or '',
        '{{data_rielaborazione}}': us.data_rielaborazione.strftime('%d/%m/%Y') if us.data_rielaborazione else '',
        '{{resp_rielaborazione}}': us.responsabile_rielaborazione or ''
    }

    # Sostituisci in tutto il documento
    _replace_placeholders_in_doc(doc, replacements)

    # Salva in BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def compile_usm_template(usm: UnitaStratigraficaMuraria, template_path: Path) -> bytes:
    """
    Compila template Word con dati USM
    Sostituisce placeholder {{...}} con dati database
    """
    # Carica template
    doc = Document(template_path)

    # Prepara dati per sostituzione
    replacements = {
        '{{usm_code}}': usm.usm_code or '',
        '{{ente_responsabile}}': usm.ente_responsabile or '',
        '{{anno}}': str(usm.anno) if usm.anno else '',
        '{{ufficio_mic}}': usm.ufficio_mic or '',
        '{{identificativo_rif}}': usm.identificativo_rif or '',
        '{{localita}}': usm.localita or '',
        '{{area_struttura}}': usm.area_struttura or '',
        '{{ambiente_unita}}': usm.ambiente_unita_funzione or '',
        '{{saggio}}': usm.saggio or '',
        '{{posizione}}': usm.posizione or '',
        '{{settori}}': usm.settori or '',
        '{{piante}}': usm.piante_riferimenti or '',
        '{{prospetti}}': usm.prospetti_riferimenti or '',
        '{{sezioni}}': usm.sezioni_riferimenti or '',
        
        # Misure e definizione
        '{{misure}}': usm.misure or '',
        '{{superficie}}': str(usm.superficie_analizzata) if usm.superficie_analizzata else '',
        '{{definizione}}': usm.definizione or '',
        
        # Tecnica costruttiva
        '{{tecnica_costruttiva}}': usm.tecnica_costruttiva or '',
        '{{sezione_tipo}}': usm.sezione_muraria_tipo or '',
        '{{sezione_spessore}}': usm.sezione_muraria_spessore or '',
        '{{funzione_statica}}': usm.funzione_statica or '',
        '{{modulo}}': usm.modulo or '',
        '{{criteri_distinzione}}': usm.criteri_distinzione or '',
        '{{provenienza_materiali}}': usm.provenienza_materiali or '',
        '{{orientamento}}': usm.orientamento or '',
        '{{uso_primario}}': usm.uso_primario or '',
        '{{riutilizzo}}': usm.riutilizzo or '',
        '{{conservazione}}': usm.stato_conservazione or '',
        
        # Materiali - formatta JSON come testo leggibile
        '{{materiali_laterizi}}': _format_json_field(usm.materiali_laterizi),
        '{{materiali_litici}}': _format_json_field(usm.materiali_elementi_litici),
        '{{materiali_altro}}': usm.materiali_altro or '',
        
        # Legante e finiture
        '{{legante}}': _format_json_field(usm.legante),
        '{{legante_altro}}': usm.legante_altro or '',
        '{{finiture}}': usm.finiture_elementi_particolari or '',
        
        # Matrix Harris - converti JSON in testo
        '{{uguale_a}}': ', '.join(map(str, usm.sequenza_fisica.get('uguale_a', []))) if usm.sequenza_fisica else '',
        '{{si_lega_a}}': ', '.join(map(str, usm.sequenza_fisica.get('si_lega_a', []))) if usm.sequenza_fisica else '',
        '{{gli_si_appoggia}}': ', '.join(
            map(str, usm.sequenza_fisica.get('gli_si_appoggia', []))) if usm.sequenza_fisica else '',
        '{{si_appoggia_a}}': ', '.join(
            map(str, usm.sequenza_fisica.get('si_appoggia_a', []))) if usm.sequenza_fisica else '',
        '{{coperto_da}}': ', '.join(map(str, usm.sequenza_fisica.get('coperto_da', []))) if usm.sequenza_fisica else '',
        '{{copre}}': ', '.join(map(str, usm.sequenza_fisica.get('copre', []))) if usm.sequenza_fisica else '',
        '{{tagliato_da}}': ', '.join(map(str, usm.sequenza_fisica.get('tagliato_da', []))) if usm.sequenza_fisica else '',
        '{{taglia}}': ', '.join(map(str, usm.sequenza_fisica.get('taglia', []))) if usm.sequenza_fisica else '',
        '{{riempito_da}}': ', '.join(map(str, usm.sequenza_fisica.get('riempito_da', []))) if usm.sequenza_fisica else '',
        '{{riempie}}': ', '.join(map(str, usm.sequenza_fisica.get('riempie', []))) if usm.sequenza_fisica else '',
        
        # Descrizione e interpretazione
        '{{descrizione}}': usm.descrizione or '',
        '{{osservazioni}}': usm.osservazioni or '',
        '{{interpretazione}}': usm.interpretazione or '',
        
        # Datazione
        '{{datazione}}': usm.datazione or '',
        '{{periodo}}': usm.periodo or '',
        '{{fase}}': usm.fase or '',
        '{{elementi_datanti}}': usm.elementi_datanti or '',
        
        # Checkbox campionature - simboli
        '{{campione_litici}}': '☑' if (usm.campionature and usm.campionature.get('elementi_litici')) else '☐',
        '{{campione_laterizi}}': '☑' if (usm.campionature and usm.campionature.get('laterizi')) else '☐',
        '{{campione_malta}}': '☑' if (usm.campionature and usm.campionature.get('malta')) else '☐',
        
        # Affidabilità e responsabilità
        '{{affidabilita}}': usm.affidabilita_stratigrafica or '',
        '{{resp_scientifico}}': usm.responsabile_scientifico or '',
        '{{data_rilevamento}}': usm.data_rilevamento.strftime('%d/%m/%Y') if usm.data_rilevamento else '',
        '{{resp_compilazione}}': usm.responsabile_compilazione or '',
        '{{data_rielaborazione}}': usm.data_rielaborazione.strftime('%d/%m/%Y') if usm.data_rielaborazione else '',
        '{{resp_rielaborazione}}': usm.responsabile_rielaborazione or ''
    }

    # Sostituisci in tutto il documento
    _replace_placeholders_in_doc(doc, replacements)

    # Salva in BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _format_json_field(json_data: Optional[Dict]) -> str:
    """Formatta campo JSON come testo leggibile"""
    if not json_data:
        return ''
    
    parts = []
    for key, value in json_data.items():
        if isinstance(value, list):
            parts.append(f"{key}: {', '.join(map(str, value))}")
        else:
            parts.append(f"{key}: {value}")
    
    return '; '.join(parts)


def _replace_placeholders_in_doc(doc: Document, replacements: Dict[str, str]):
    """
    Sostituisce placeholder in tutto il documento (tabelle, paragrafi, header, footer)
    Mantiene formattazione originale
    """
    # Sostituisci in paragrafi
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Sostituisci in tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Sostituisci in header/footer
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """Sostituisce placeholder in paragrafo mantenendo formattazione"""
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Sostituisci in tutti i run
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))


def generate_us_filename(us: UnitaStratigrafica) -> str:
    """Genera nome file standard per US"""
    us_number_str = us.us_code.replace('US', '').replace('us', '').lstrip('0') if us.us_code else '0'
    try:
        us_number_int = int(us_number_str)
    except ValueError:
        us_number_int = 0

    localita_clean = (us.localita or '').replace(' ', '_').replace(',', '') or 'Sito'
    anno = us.anno or datetime.now().year

    return f"US{us_number_int:03d}_{localita_clean}_{anno}.docx"


def generate_usm_filename(usm: UnitaStratigraficaMuraria) -> str:
    """Genera nome file standard per USM"""
    usm_number_str = usm.usm_code.replace('USM', '').replace('usm', '').lstrip('0') if usm.usm_code else '0'
    try:
        usm_number_int = int(usm_number_str)
    except ValueError:
        usm_number_int = 0

    localita_clean = (usm.localita or '').replace(' ', '_').replace(',', '') or 'Sito'
    anno = usm.anno or datetime.now().year

    return f"USM{usm_number_int:03d}_{localita_clean}_{anno}.docx"


# ===== ENDPOINT EXPORT SINGOLA US =====

@router.get("/us/{us_id}/word")
async def export_us_word(
        us_id: UUID,
        db: AsyncSession = Depends(get_async_session),
        current_user_id: UUID = Depends(get_current_user_id_with_blacklist),
        user_sites: List[Dict[str, Any]] = Depends(get_current_user_sites_with_blacklist)
):
    """
    Esporta singola scheda US in formato Word
    Usa template con placeholder - mantiene layout identico al MiC 2021
    """
    try:
        logger.info(f"→ Export Word US {us_id}")

        # Carica US
        us_query = select(UnitaStratigrafica).where(UnitaStratigrafica.id == us_id)
        us_result = await db.execute(us_query)
        us = us_result.scalar_one_or_none()

        if not us:
            raise HTTPException(status_code=404, detail="US non trovata")

        logger.debug(f"US trovata: {us.us_code}, sito: {us.site_id}")

        # Verifica accesso
        if not await verify_site_access(us.site_id, user_sites):
            raise HTTPException(status_code=403, detail="Accesso negato al sito")

        # Compila template
        doc_bytes = compile_us_template(us, US_TEMPLATE_PATH)
        filename = generate_us_filename(us)

        logger.info(f"✓ Export completato: {filename} ({len(doc_bytes)} bytes)")

        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Errore export Word US {us_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Errore generazione documento: {str(e)}")


# ===== ENDPOINT EXPORT MULTIPLO US =====

@router.post("/us/bulk-word")
async def export_multiple_us_word(
        us_ids: List[UUID],
        db: AsyncSession = Depends(get_async_session),
        current_user_id: UUID = Depends(get_current_user_id_with_blacklist),
        user_sites: List[Dict[str, Any]] = Depends(get_current_user_sites_with_blacklist)
):
    """
    Esporta multiple US come ZIP di documenti Word
    Ogni US genera un file separato con template compilato
    """
    if len(us_ids) > 50:
        raise HTTPException(status_code=400, detail="Troppi US selezionati (max 50)")

    try:
        logger.info(f"→ Export bulk {len(us_ids)} US")

        # Carica tutte le US
        us_query = select(UnitaStratigrafica).where(UnitaStratigrafica.id.in_(us_ids))
        us_result = await db.execute(us_query)
        us_list = us_result.scalars().all()

        if not us_list:
            raise HTTPException(status_code=404, detail="Nessuna US trovata")

        # Verifica accessi
        for us in us_list:
            if not await verify_site_access(us.site_id, user_sites):
                raise HTTPException(status_code=403, detail=f"Accesso negato per US {us.us_code}")

        # Crea ZIP
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            success_count = 0

            for us in us_list:
                try:
                    # Compila template
                    doc_bytes = compile_us_template(us, US_TEMPLATE_PATH)
                    filename = generate_us_filename(us)

                    # Aggiungi a ZIP
                    zip_file.writestr(filename, doc_bytes)
                    success_count += 1
                    logger.debug(f"US {us.us_code} aggiunta al ZIP")

                except Exception as e:
                    logger.error(f"Errore generazione US {us.us_code}: {str(e)}")
                    continue

        zip_buffer.seek(0)

        # Nome ZIP
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_filename = f"SchediUS_Export_{timestamp}.zip"

        logger.info(f"✓ Export bulk completato: {success_count}/{len(us_list)} US")

        return StreamingResponse(
            io.BytesIO(zip_buffer.read()),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={zip_filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Errore export bulk: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Errore export multiplo: {str(e)}")


# ===== ENDPOINT EXPORT SITO COMPLETO =====

@router.get("/site/{site_id}/us/word-zip")
async def export_site_us_word(
        site_id: UUID,
        validated_only: bool = False,
        db: AsyncSession = Depends(get_async_session),
        current_user_id: UUID = Depends(get_current_user_id_with_blacklist),
        user_sites: List[Dict[str, Any]] = Depends(get_current_user_sites_with_blacklist)
):
    """
    Esporta tutte le US di un sito come ZIP
    Opzione per esportare solo US validate
    """
    try:
        logger.info(f"→ Export sito {site_id} (validated_only={validated_only})")

        # Verifica accesso
        if not await verify_site_access(site_id, user_sites):
            raise HTTPException(status_code=403, detail="Accesso negato al sito")

        # Query US del sito
        us_query = select(UnitaStratigrafica).where(UnitaStratigrafica.site_id == site_id)

        if validated_only:
            us_query = us_query.where(UnitaStratigrafica.is_validated == True)

        us_query = us_query.order_by(UnitaStratigrafica.us_code)
        us_result = await db.execute(us_query)
        us_list = us_result.scalars().all()

        if not us_list:
            raise HTTPException(status_code=404, detail="Nessuna US trovata per questo sito")

        if len(us_list) > 200:
            raise HTTPException(status_code=400, detail=f"Troppe US nel sito ({len(us_list)}). Max 200.")

        # Genera ZIP
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            success_count = 0

            for us in us_list:
                try:
                    doc_bytes = compile_us_template(us, US_TEMPLATE_PATH)
                    filename = generate_us_filename(us)
                    zip_file.writestr(filename, doc_bytes)
                    success_count += 1
                except Exception as e:
                    logger.error(f"Errore US {us.us_code}: {str(e)}")
                    continue

        zip_buffer.seek(0)

        # Nome ZIP
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        validated_suffix = "_Validate" if validated_only else "_Tutte"
        zip_filename = f"SchediUS_Sito{validated_suffix}_{timestamp}.zip"

        logger.info(f"✓ Export sito completato: {success_count}/{len(us_list)} US")

        return StreamingResponse(
            io.BytesIO(zip_buffer.read()),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={zip_filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Errore export sito {site_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Errore export sito: {str(e)}")


# ===== ENDPOINT SINGOLA USM =====

@router.get("/usm/{usm_id}/word")
async def export_usm_word(
        usm_id: UUID,
        db: AsyncSession = Depends(get_async_session),
        current_user_id: UUID = Depends(get_current_user_id_with_blacklist),
        user_sites: List[Dict[str, Any]] = Depends(get_current_user_sites_with_blacklist)
):
    """Esporta singola scheda USM in formato Word"""
    try:
        logger.info(f"→ Export Word USM {usm_id}")

        # Carica USM
        usm_query = select(UnitaStratigraficaMuraria).where(UnitaStratigraficaMuraria.id == usm_id)
        usm_result = await db.execute(usm_query)
        usm = usm_result.scalar_one_or_none()

        if not usm:
            raise HTTPException(status_code=404, detail="USM non trovata")

        logger.debug(f"USM trovata: {usm.usm_code}, sito: {usm.site_id}")

        # Verifica accesso
        if not await verify_site_access(usm.site_id, user_sites):
            raise HTTPException(status_code=403, detail="Accesso negato al sito")

        # Compila template
        doc_bytes = compile_usm_template(usm, USM_TEMPLATE_PATH)
        filename = generate_usm_filename(usm)

        logger.info(f"✓ Export USM completato: {filename} ({len(doc_bytes)} bytes)")

        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Errore export USM {usm_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Errore generazione USM: {str(e)}")


# ===== UTILITY ENDPOINTS =====

@router.get("/supported-formats")
async def get_supported_export_formats():
    """Formati export supportati"""
    return {
        'formats': {
            'word': {
                'extension': '.docx',
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'description': 'Documento Word da template con placeholder',
                'template_based': True,
                'template_file': 'US_Template_con_Placeholder.docx'
            }
        },
        'export_options': {
            'single_us': '/api/us-export/us/{id}/word',
            'single_usm': '/api/us-export/usm/{id}/word',
            'bulk_us': '/api/us-export/us/bulk-word',
            'site_export': '/api/us-export/site/{site_id}/us/word-zip'
        },
        'template_info': {
            'method': 'Template con placeholder {{...}}',
            'compliance': 'MiC Standard 2021',
            'layout': 'Identico al documento originale',
            'placeholders_count': 48
        }
    }
